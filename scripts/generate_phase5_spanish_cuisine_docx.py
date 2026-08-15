from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from generate_book_blueprint_docx import configure_document, add_markdown_document


OUTPUT = Path(
    "attached_assets/mastering_french_spanish_cooking_phase5_spanish_cuisine.docx"
)


def add_centered_text(doc, text, size=12, bold=False, italic=False, color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_title_page(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(105)
    run = paragraph.add_run("MASTERING THE ART OF\nFRENCH & SPANISH COOKING")
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(107, 62, 38)

    add_centered_text(
        doc,
        "A Regional, Technique-Driven Guide to the Ingredients, "
        "Traditions, and Recipes of France and Spain",
        size=14,
        italic=True,
        color="8A5A3B",
    )
    add_centered_text(doc, "PHASE 5 — SPANISH CUISINE", size=12, bold=True)
    add_centered_text(doc, "By [AUTHOR NAME]", size=12)
    doc.add_page_break()


SPANISH_CUISINE = r"""
# PHASE 5 — SPANISH CUISINE

## A note to the cook

Spanish cooking is not one flavor profile and not one national menu. It is a conversation among Atlantic coasts, Mediterranean harbors, dry inland plains, green northern hills, island kitchens, mountain villages, and cities shaped by trade and migration. A Galician cook, a Valencian rice cook, a Basque grill cook, and an Andalusian gazpacho maker may share olive oil and garlic while using them in entirely different ways.

This section approaches Spain region by region and technique by technique. The aim is not to turn every dish into a museum piece. It is to explain the choices that give Spanish food its character: the slow cooking of a sofrito, the controlled bitterness of good olive oil, the perfume of saffron, the smoke of pimentón, the contrast between crisp and soft, the use of preserved ingredients, and the generosity of food placed in the middle of the table.

Regional names in this section are guides, not rigid borders. Dishes travel. Ingredients cross provincial lines. Families alter recipes according to season, budget, religion, available equipment, and memory. When a preparation has several established forms, the recipe identifies the style used and explains what should remain constant.

## The Spanish kitchen in practice

Spanish food rewards preparation before heat. Chop the onion and pepper, measure the rice, warm the stock, dry the seafood, and put the serving dishes within reach. Once a pan is hot, many decisions must happen quickly. A good mise en place leaves the cook free to notice the important cues: whether the sofrito has lost its raw edge, whether the rice is absorbing evenly, whether the oil is hot enough for a crisp fritter, and whether a shellfish has opened.

Extra-virgin olive oil is used throughout this section, but its flavor and heat behavior vary. Choose a sound oil and taste it. Use a milder oil when delicate ingredients should lead and a more assertive oil when bitterness and fruitiness are part of the design. Salted fish, cured meats, olives, stock, and preserved peppers can all contribute salt, so season in stages and taste before finishing.

Spanish meals are often social by design. A tapa can be a small serving, a snack, or a way of eating; a ración is a larger portion for sharing; a pincho is commonly assembled or served on bread or a skewer. The exact vocabulary changes by place. What matters in the home kitchen is to offer contrast: hot and cold, crisp and tender, rich and bright, quick and slow.

# CHAPTER 18 — UNDERSTANDING SPANISH CUISINE

## A cuisine of landscapes

Spain’s culinary geography is unusually varied. The Atlantic north has cool water, green produce, cider, beans, shellfish, and dishes that make good use of gentle braising. The Mediterranean coast brings olive oil, rice, citrus, vegetables, almonds, and seafood. The central plateau has dry climates, wheat, pulses, lamb, pork, game, and stews that conserve heat and flavor. Andalusia contributes cold soups, fried foods, sherry, almonds, and the deep agricultural history of the south. The islands add their own produce, preserved foods, and oceanic connections.

Climate affects not only ingredients but also the form of a meal. A chilled soup makes sense in a hot inland summer. A pot of legumes can feed a household through a cold winter. Small plates make it possible to share several tastes, while rice dishes turn a modest set of ingredients into a communal centerpiece.

## Layers of influence

Spanish food reflects long histories of exchange. Roman agriculture, Jewish and Muslim foodways, Atlantic trade, American crops, pastoral traditions, monastic kitchens, fishing communities, and modern urban cooking all form part of the story. Tomatoes, peppers, potatoes, and chocolate became central to later Spanish cooking only after their arrival from the Americas. Rice and techniques for irrigation and flavoring developed through Mediterranean and Islamic connections.

Historical influence should not be reduced to a single origin story. A dish can have a documented lineage, a family legend, and a modern form at the same time. This book distinguishes what is broadly established from what is a useful culinary interpretation.

## The regions in the kitchen

**Andalusia** is associated with olive oil, sherry vinegar, almonds, cold soups, fried fish, citrus, and the use of bread to give body. **Catalonia** makes broad use of sofregit, beans, seafood, mushrooms, picada, and sauces built in a mortar. **Valencia** is strongly associated with rice cookery, beans, green vegetables, saffron, and the social ritual of cooking in a wide pan.

**Galicia** brings exceptional seafood, potatoes, peppers, pork, and broths. **The Basque Country** is known for careful handling of fish, peppers, stews, cider, and a respect for excellent ingredients. **Asturias** is home to cider, beans, cabbage, pork, and hearty northern cooking. **Madrid and the central regions** contribute cocidos, roast meats, legumes, and sauces suited to a continental climate. Castilla y León, Castilla-La Mancha, Extremadura, Murcia, Aragón, and the Canary Islands each add distinct products and methods.

These labels help a cook ask better questions: Is the dish meant to be fresh or slow-cooked? Is the main seasoning smoke, acid, saffron, garlic, or the taste of the ingredient itself? Is it served alone, with bread, or as part of a spread?

## A working philosophy

Four principles recur:

- **Start with a flavor base.** Onion, garlic, pepper, tomato, olive oil, and pimentón may be combined in different proportions, but slow aromatic cooking is often the first layer.
- **Protect the main ingredient.** Seafood is not improved by overcooking, and a ripe tomato does not need to be buried under spices.
- **Use contrast deliberately.** Crisp potatoes meet a soft sauce; cold soup meets a garnish; a rich croqueta meets a sharp pickle or olive.
- **Cook for sharing.** A dish should taste complete on its own but also fit beside other foods on the table.

Spanish authenticity is not a demand for inaccessible ingredients. It is attention to the defining method and balance. If saffron is replaced, the result should be described honestly; if a paella pan is unavailable, the cook should choose the widest suitable pan and accept that the rice layer will change.

**[PHOTO SUGGESTION: Regional Spanish pantry — olive oil, saffron, pimentón, almonds, olives, dried beans, rice, garlic, peppers, and tomatoes arranged by texture and color.]**

# CHAPTER 19 — ESSENTIAL SPANISH INGREDIENTS

## Olive oil

Extra-virgin olive oil is a seasoning, cooking medium, finishing ingredient, and carrier of aroma. Fruitiness, bitterness, and pepperiness are signs of character rather than defects, although the oil should not taste rancid or stale. Use a balanced oil for everyday cooking and a more expressive oil for finishing bread, soups, salads, and grilled vegetables.

Store oil sealed in a cool, dark place. Heat changes its aroma, so a delicate finishing oil may be better kept for the table. Olive oil can be substituted with another neutral oil in a technical emergency, but the dish will lose an important part of its Spanish identity.

## Garlic, onions, and peppers

Garlic may be sliced, crushed, or made into a paste. Thin slices brown quickly and are ideal for a short pan sauce; crushed cloves perfume a stew; a paste distributes flavor through a marinade or alioli. Onions are often cooked patiently until sweet. Red and green peppers contribute color, sweetness, and sometimes heat. Dried ñora or similar peppers may be rehydrated and scraped into sauces where available.

## Tomatoes

Tomatoes may be grated, crushed, roasted, dried, or cooked down. Grating a ripe tomato leaves the skin behind and produces a quick pulp for pan con tomate. Canned tomatoes are useful when fresh tomatoes lack flavor. In a sofrito, the tomato is not simply added and boiled; it is reduced until the oil begins to show and the raw watery edge disappears.

## Pimentón and saffron

Pimentón is ground dried pepper. Sweet, bittersweet, and hot versions exist, and smoked pimentón adds a wood-fire note. Add it briefly to warm oil or a moist base, then protect it from burning. Burnt paprika becomes harsh.

Saffron contributes fragrance, color, and a subtle earthy bitterness. Use a small amount, crush it gently, and bloom it in warm liquid when the recipe permits. Turmeric may color a dish, but it is not an identical substitute for saffron; label the adaptation honestly.

## Vinegar, olives, almonds, and preserved foods

Sherry vinegar is assertive and complex, while wine vinegar and cider vinegar are milder alternatives. Add vinegar gradually to gazpacho, salads, marinades, and sauces. Olives, capers, anchovies, salted cod, cured sausage, and preserved peppers concentrate flavor and salt. Rinse only when the recipe calls for a cleaner taste; rinsing can remove useful character.

Almonds and other nuts add body to picada, cold soups, desserts, and sauces. Toast them gently and watch closely. Their color can deepen quickly after the aroma becomes fragrant.

## Rice and legumes

Rice is not interchangeable by name alone. Short- and medium-grain varieties absorb stock and release some starch while remaining distinct when treated correctly. Long-grain rice behaves differently and may not produce the texture expected in a regional rice dish. Measure the pan, liquid, and rice before starting.

Chickpeas, beans, and lentils are central to many Spanish stews. Dried legumes reward soaking and gentle cooking; canned legumes are practical for weeknight meals. Salt and acid do not always prevent tenderness, but older beans, hard water, and premature acid can all lengthen cooking. Cook until tender, then adjust the final seasoning.

## Seafood and cured meats

Spanish seafood ranges from tiny fish for frying to octopus, shellfish, salt cod, squid, and large Atlantic fish. Buy from a reputable source, keep it cold, and use clean boards and utensils. Shellfish should smell clean and briny; discard specimens with broken shells or those that do not close when handled when local guidance recommends it.

Chorizo, jamón, cured pork, and other preserved meats are seasonings as much as proteins. Taste before salting. When a recipe uses a local product unavailable elsewhere, choose the closest functional alternative and explain the difference.

# CHAPTER 20 — TAPAS: THE HEART OF SPANISH DINING

## Small plates, large principles

Tapas are not defined only by small size. They are part of a way of eating in which conversation, movement, sharing, and variety matter. A tapas table should not contain nine dishes that are all fried or all tomato-based. Build a spread with a rhythm: one crisp item, one cool vegetable dish, one substantial plate, one seafood preparation, something pickled or acidic, and bread.

Make-ahead elements are valuable. Marinated olives, roasted peppers, sauces, and cooked beans can be prepared first. Frying, grilling, and final garnishes should happen close to service. Serve in small portions, but do not turn every dish into a precious bite; a bowl of potatoes or a plate of tortilla should feel generous.

## Recipe: Patatas bravas

### Recipe information

- **Prep Time:** 20 minutes
- **Cook Time:** 35–45 minutes
- **Total Time:** About 1 hour
- **Servings:** 4 as a tapa
- **Difficulty:** Intermediate
- **Cuisine:** Spanish
- **Region:** Madrid-style

### Equipment

Heavy frying pan or fryer, saucepan, thermometer if frying, paper-lined tray.

### Ingredients

- 800 g (1¾ lb) waxy potatoes, peeled and cut into 3 cm (1¼ in) pieces
- 750 ml (3 cups) frying oil
- 1 tablespoon olive oil
- 1 small onion, finely chopped
- 2 garlic cloves, minced
- 1 teaspoon smoked pimentón
- ½ teaspoon hot pimentón or cayenne
- 250 g (9 oz) crushed tomatoes
- 1 tablespoon sherry vinegar
- Salt

### Instructions

1. Cover the potatoes with cold water, bring to a gentle simmer, and cook for 6–8 minutes until the edges soften but the centers remain firm. Drain and steam-dry.
2. For the sauce, warm the olive oil over medium-low heat. Cook the onion for 8–10 minutes until soft. Add garlic and pimentón for 30 seconds, then add tomatoes. Simmer 15 minutes, blend if desired, and finish with vinegar and salt.
3. Heat the frying oil to 150°C (300°F). Fry the potatoes in batches for 5–6 minutes until tender but pale. Drain.
4. Raise the oil to 185°C (365°F). Fry again until deeply golden and crisp. Drain, salt immediately, and serve with warm brava sauce.

### Chef’s tips, mistakes, and storage

Dry potatoes fry more safely and become crisper. Overcrowding lowers the oil temperature and produces greasy potatoes. The sauce may be made up to 3 days ahead and refrigerated; fry the potatoes just before serving.

### Cultural note

The name describes the assertive sauce, but formulas vary widely. Some versions are smooth, some chunky, and some are served with alioli as well as brava sauce.

## Recipe: Tortilla española

### Recipe information

- **Prep Time:** 15 minutes
- **Cook Time:** 30 minutes
- **Total Time:** 45 minutes
- **Servings:** 4–6
- **Difficulty:** Intermediate
- **Cuisine:** Spanish
- **Region:** Widely prepared across Spain

### Ingredients

- 700 g (1½ lb) waxy potatoes, thinly sliced
- 1 medium onion, thinly sliced
- 250 ml (1 cup) olive oil
- 6 large eggs
- 1½ teaspoons salt, divided

### Instructions

1. Warm the oil in a 24 cm (9½ in) nonstick skillet over medium-low heat. Add the potatoes and onion with 1 teaspoon salt. Cook gently for 15–20 minutes, turning occasionally, until tender but not browned.
2. Drain the mixture, reserving 2 tablespoons oil. Beat the eggs with the remaining salt. Fold in the warm potatoes and rest for 5 minutes.
3. Heat the reserved oil in the skillet. Add the egg mixture and cook over medium-low heat until the edges set and the center remains loose.
4. Place a plate over the skillet, invert decisively, and slide the tortilla back into the pan. Cook 2–4 minutes more, depending on the desired center. Rest 5 minutes before slicing.

### Chef’s tips and variations

The potatoes should soften in oil rather than crisp like chips. A runny center is a matter of preference, but eggs should be handled with appropriate food-safety care. Add roasted pepper, herbs, or a small amount of chorizo only when the variation is intentional.

## Recipe: Gambas al ajillo

### Recipe information

- **Prep Time:** 10 minutes
- **Cook Time:** 5 minutes
- **Total Time:** 15 minutes
- **Servings:** 4 as a tapa
- **Difficulty:** Beginner
- **Cuisine:** Spanish
- **Region:** Common throughout Spain

### Ingredients

- 500 g (1 lb 2 oz) peeled raw shrimp, patted dry
- 90 ml (6 tablespoons) olive oil
- 5 garlic cloves, thinly sliced
- 1 small dried chile, crumbled, optional
- 1 tablespoon chopped parsley
- ½ teaspoon salt
- Crusty bread, to serve

### Instructions

1. Warm the oil, garlic, and chile in a small, wide pan over medium-low heat until the garlic is pale gold and fragrant.
2. Raise the heat to medium-high. Add the shrimp in one layer and cook 60–90 seconds per side, just until opaque and curled.
3. Remove from heat, add parsley and salt, and serve immediately with bread for the infused oil.

### Common mistakes

Black garlic will make the oil bitter. Cold, wet shrimp will steam instead of sear. The oil continues cooking the shrimp after the pan leaves the heat, so stop as soon as the thickest piece is opaque.

## More tapas foundations

**Croquetas de patata y jamón** teach a thick béchamel-like filling, chilling, breading, and controlled frying. **Pimientos de Padrón** are blistered quickly in olive oil and finished with flaky salt. **Albóndigas** begin with a moist meat mixture and finish in a tomato, almond, or wine sauce. **Champiñones al ajillo** demonstrate that mushrooms need space and high enough heat to lose their water. **Verduras a la plancha** succeed when vegetables are dry, lightly oiled, and not turned too often. **Aceitunas aliñadas** show how citrus peel, herbs, garlic, and vinegar can transform a preserved ingredient without hiding it.

For a balanced tapas evening, prepare the olives and sauces first, the tortilla next, the vegetables and mushrooms shortly before serving, and the fried items last.

**[PHOTO SUGGESTION: Tapas spread — tortilla wedges, brava potatoes, garlic shrimp, blistered peppers, marinated olives, and bread arranged for sharing.]**

# CHAPTER 21 — SPANISH SOUPS AND SALADS

## Heat and coolness

Spanish soups range from substantial bean pots to chilled vegetable purées. A cold soup must be seasoned more assertively than a warm one because chilling dulls aroma and salt perception. Oil, vinegar, and water need to be balanced rather than added by rote.

## Recipe: Gazpacho andaluz

### Recipe information

- **Prep Time:** 20 minutes plus chilling
- **Cook Time:** None
- **Total Time:** About 2 hours
- **Servings:** 6
- **Difficulty:** Beginner
- **Cuisine:** Andalusian

### Ingredients

- 1 kg (2¼ lb) ripe tomatoes, cored and chopped
- 1 small cucumber, peeled and chopped
- 1 red bell pepper, chopped
- 1 small garlic clove
- 100 g (3½ oz) day-old bread, torn
- 80 ml (⅓ cup) extra-virgin olive oil
- 45 ml (3 tablespoons) sherry vinegar
- 200 ml (¾ cup plus 1 tablespoon) cold water, plus more as needed
- 1½ teaspoons salt

### Instructions

1. Combine the tomatoes, cucumber, pepper, garlic, bread, vinegar, and salt. Rest 20 minutes so the bread softens and the vegetables release juice.
2. Blend until smooth, slowly streaming in the oil. Add water to reach a pourable consistency.
3. Taste for salt and vinegar, strain for a refined texture if desired, and chill at least 1 hour. Stir before serving.

### Chef’s tips

The soup should taste bright, cool, and clearly of tomato. Over-blending can warm the mixture and make it foamy; chill it again before service. Garnish with diced vegetables, croutons, or a thread of olive oil.

## Recipe: Salmorejo cordobés

Salmorejo is thicker and more bread-forward than gazpacho. Blend 1 kg (2¼ lb) ripe tomatoes with 200 g (7 oz) day-old white bread, 1 small garlic clove, 1½ teaspoons salt, and 30 ml (2 tablespoons) sherry vinegar. Stream in 120 ml (½ cup) olive oil until glossy. Chill and garnish with chopped hard-boiled egg and finely sliced cured ham. Add the oil gradually; a rushed addition can make the mixture feel greasy rather than creamy.

## Recipe: Ajo blanco

Soak 150 g (5¼ oz) blanched almonds and 100 g (3½ oz) stale bread in 500 ml (2 cups) cold water for 20 minutes. Blend with 1 garlic clove, 30 ml (2 tablespoons) sherry vinegar, and 1 teaspoon salt. Stream in 100 ml (scant ½ cup) olive oil, chill, and thin to taste. Serve with grapes or melon. The soup should be nutty, pale, and refreshing, not aggressively garlicky.

## Warm soups and salads

A Spanish bean soup begins with a sofrito or a cured-meat base, then adds soaked beans and enough liquid for a gentle simmer. Add greens when the beans are nearly tender so they retain color and character. Lentil soup is more forgiving and can be built with carrot, onion, tomato, pimentón, and olive oil; avoid boiling hard, which can split the skins.

For a tomato-and-onion salad, choose ripe tomatoes, salt them shortly before serving, and dress with olive oil and vinegar. An orange-and-olive salad benefits from contrast: sweet citrus, briny olives, thin onion, toasted fennel or cumin, and enough oil to connect the flavors.

# CHAPTER 22 — PAELLA AND SPANISH RICE DISHES

## Rice is a method, not a garnish

Rice dishes depend on geometry. A wide pan creates a shallow layer in which liquid evaporates evenly and grains cook without excessive stirring. A narrow pot creates a deeper layer and a different texture. Weigh the rice and measure the stock before beginning. Once the rice is distributed, stirring is usually minimized so the grains remain distinct.

## Sofrito, stock, and heat

Build the flavor base patiently. Add rice only when the aromatics are concentrated and the liquid is ready. Stock should be well seasoned but not aggressively salty because evaporation concentrates it. Saffron may be steeped in warm stock. Maintain an active simmer early, then lower the heat as the liquid is absorbed.

Socarrat is a prized toasted layer at the bottom of some rice dishes. It requires enough heat and a cook’s attention. A faint crackle and roasted aroma can signal development; a sharp burnt smell means the pan must leave the heat immediately.

## Recipe: Paella Valenciana-style

### Recipe information

- **Prep Time:** 30 minutes
- **Cook Time:** 35–40 minutes
- **Total Time:** About 1 hour 15 minutes
- **Servings:** 4–6
- **Difficulty:** Advanced
- **Cuisine:** Spanish
- **Region:** Valencia-inspired

### Ingredients

- 400 g (14 oz) short- or medium-grain paella rice
- 700 g (1½ lb) chicken thighs and rabbit or additional chicken, cut into pieces
- 150 g (5 oz) flat green beans, cut into pieces
- 100 g (3½ oz) cooked white beans or butter beans
- 2 ripe tomatoes, grated
- 1 red pepper, diced
- 3 garlic cloves, minced
- 60 ml (¼ cup) olive oil
- 1 teaspoon sweet pimentón
- A pinch of saffron threads
- 1.2 litres (5 cups) hot chicken stock
- Salt, lemon wedges, and rosemary to serve

### Instructions

1. Heat the oil in a 34–36 cm (13–14 in) paella pan. Salt the chicken and brown it thoroughly on all sides. Move it toward the edge.
2. Add the beans and pepper. Cook until lightly colored. Add the garlic, tomato, and pimentón; cook until the tomato is concentrated and the oil shows at the edges.
3. Add the rice and stir for 1 minute to coat. Pour in the hot stock and add the saffron. Distribute the rice evenly; after this point, avoid stirring.
4. Simmer actively for 8 minutes, then lower the heat and cook 10–12 minutes more. Add the white beans during the final 8 minutes.
5. When the liquid is nearly absorbed and the rice is tender with a slight bite, increase the heat briefly to develop socarrat. Rest off the heat for 5–10 minutes before serving.

### Chef’s tips and variations

The pan should fit the burner; rotate it over a small burner or use a diffuser if needed. A crowded pan produces uneven rice. Seafood, vegetable, and mixed versions require different timing, so do not simply exchange ingredients without adjusting the stock and cooking stages.

## Recipe: Seafood paella

For 4 servings, cook a concentrated sofrito from 1 onion, 1 red pepper, 3 garlic cloves, 2 grated tomatoes, 60 ml olive oil, and 1 teaspoon pimentón. Add 350 g (12 oz) rice, 1 litre (4¼ cups) seafood broth, and saffron. Simmer without stirring. Add squid with the sofrito, then place mussels, clams, and shrimp during the last 8–10 minutes so they cook without becoming tough. Discard shells that remain closed after cooking according to local food-safety guidance. Rest before serving.

## Other Spanish rice dishes

**Arroz negro** uses squid or cuttlefish and its ink for a dark, marine flavor, usually with alioli. **Arroz al horno** is baked in a separate vessel and can include chickpeas, pork, tomato, and potato. **Vegetable rice** should treat vegetables according to their water content: firm beans and artichokes enter early, tender greens late. **Catalan-style rice** may be more brothy or finished with picada, depending on the preparation. The word “paella” should not be used to erase these distinctions.

**[PHOTO SUGGESTION: Step-by-step rice sequence — concentrated sofrito, rice and stock spread in the pan, final absorption, and a close view of a lightly developed socarrat.]**

# CHAPTER 23 — SPANISH SEAFOOD

## Buy well, cook briefly

Seafood reveals poor technique quickly. Keep it cold, dry its surface before searing, and season according to the salinity of the ingredient. Shellfish release liquid as they open; do not drown them in a pan of cold sauce. A fish fillet is ready when the thickest part turns opaque and flakes with gentle pressure, although the exact cue depends on species and thickness.

## Recipe: Pulpo a la gallega

### Recipe information

- **Prep Time:** 20 minutes
- **Cook Time:** 60–90 minutes plus resting
- **Total Time:** About 2 hours
- **Servings:** 4
- **Difficulty:** Intermediate
- **Cuisine:** Galician

### Ingredients

- 1.5 kg (3¼ lb) cleaned octopus
- 2 bay leaves
- 1 onion, halved
- 1 kg (2¼ lb) waxy potatoes
- Extra-virgin olive oil
- Sweet pimentón and hot pimentón
- Coarse salt

### Instructions

1. Bring a large pot of water with the bay and onion to a gentle boil. Hold the octopus by the head and dip the tentacles into the water three times, allowing them to curl, then lower it fully.
2. Simmer gently, partially covered, until a knife enters the thickest tentacle with little resistance, usually 60–90 minutes depending on size. Rest it in the cooking liquid for 15 minutes.
3. Boil the potatoes separately until tender. Slice them onto a wooden board or warm platter.
4. Cut the octopus into pieces and arrange over the potatoes. Finish with olive oil, pimentón, and coarse salt.

### Cultural note

The familiar presentation is associated with Galician fairs and communal service. The preparation is simple, but texture depends on gentle cooking and adequate resting.

## Recipe: Squid with onions

Cook 700 g (1½ lb) cleaned squid in a wide pan over high heat in 30 ml olive oil until it releases and then reabsorbs its liquid; remove. Lower the heat, cook 2 sliced onions with a pinch of salt for 25 minutes until golden, add 2 garlic cloves and 100 ml dry white wine, and reduce. Return the squid for 2–3 minutes, finish with parsley, and serve immediately. Long cooking can tenderize squid, but an in-between stage often makes it rubbery.

## Recipe: Clams in sherry-style sauce

Rinse 1 kg (2¼ lb) live clams and discard cracked shells. Sweat 2 minced garlic cloves in 30 ml olive oil, add 100 ml dry sherry and 100 ml water, then add the clams. Cover and cook until open. Stir in parsley and a small amount of flour-thickened stock only if a sauce with body is desired. Do not force closed shells open.

## More coastal preparations

Steamed mussels need only a short, hot cooking time and a lid that traps steam. Salt cod must be desalted in several changes of cold water, kept refrigerated, and tasted before cooking. Sardines benefit from a hot grill or pan, dry skin, and a sharp finish of lemon, parsley, or vinegar. **Merluza en salsa verde** uses gentle poaching and parsley rather than heavy browning. **Gambas a la plancha** rely on a very hot surface and a short cooking window.

# CHAPTER 24 — SPANISH MEAT AND POULTRY

## Browning, smoke, and slow cooking

Spanish meat cookery moves between fast garlic-scented sautéing and long, patient stews. Dry meat before browning. Use pimentón with care. When a braise includes tomato, wine, stock, or paprika, reduce the liquid enough to concentrate it before covering the pot.

## Recipe: Pollo al ajillo

### Recipe information

- **Prep Time:** 15 minutes
- **Cook Time:** 35 minutes
- **Total Time:** 50 minutes
- **Servings:** 4
- **Difficulty:** Beginner
- **Cuisine:** Spanish

### Ingredients

- 1.2 kg (2½ lb) bone-in chicken thighs and drumsticks
- Salt and black pepper
- 60 ml (¼ cup) olive oil
- 10 garlic cloves, lightly crushed
- 150 ml (⅔ cup) dry white wine
- 1 bay leaf
- 2 tablespoons chopped parsley

### Instructions

1. Dry and season the chicken. Heat the oil in a wide pan and brown the chicken in batches.
2. Add the garlic and bay leaf. Cook gently for 2 minutes, turning the garlic so it becomes golden rather than dark.
3. Return all chicken to the pan. Add the wine, cover loosely, and simmer 25–30 minutes until the chicken is cooked through and tender.
4. Uncover and reduce the liquid until glossy. Finish with parsley and taste for salt.

### Common mistakes

Burnt garlic cannot be repaired. If it darkens before the chicken is ready, remove it and return fresh garlic near the end. The sauce should cling lightly; if it is watery, reduce it after the chicken is cooked.

## Recipe: Albóndigas in tomato sauce

Mix 500 g (1 lb 2 oz) ground beef and pork with 50 g (1¾ oz) soaked bread, 1 egg, 1 minced garlic clove, parsley, salt, and pepper. Shape into small balls, brown in olive oil, and remove. Cook onion and carrot until soft, add 400 g crushed tomato and 150 ml stock, then return the meatballs and simmer 20 minutes. Finish with a splash of sherry vinegar. The panade keeps the meat tender; compact, overworked meatballs become dense.

## Recipe: Catalan-style chicken with picada

Brown 1.2 kg chicken pieces and remove. Cook 1 chopped onion, 2 garlic cloves, and 2 grated tomatoes until concentrated. Add 150 ml white wine and 500 ml stock; return the chicken and braise covered at 170°C (340°F) for 35–45 minutes. Pound 40 g toasted almonds, 1 slice fried bread, 1 garlic clove, parsley, and a pinch of saffron into a coarse picada. Stir it into the sauce for the final 5 minutes.

## Braises and family pots

Pork with pimentón should balance smoke with acidity from wine or vinegar. Lamb stew benefits from browning, rosemary or other appropriate herbs, and a broth that is reduced before serving. Beef braises need time rather than high heat; test with a fork, not a clock alone. Cocido madrileño is a family of layered preparations involving chickpeas, meat, vegetables, and broth. Its service may separate the broth, legumes, and meats into courses.

## Recipe: Lentils with vegetables

Sweat 1 onion, 1 carrot, 1 leek, and 1 red pepper in 45 ml olive oil for 12 minutes. Add 250 g (9 oz) rinsed brown lentils, 1 grated tomato, 1 bay leaf, 1 teaspoon pimentón, and 1.2 litres stock. Simmer gently 30–40 minutes until tender. Add salt near the end and finish with sherry vinegar. Keep the lentils intact; vigorous boiling can make the broth muddy and the vegetables disappear.

# CHAPTER 25 — SPANISH VEGETABLES AND SIDE DISHES

## Vegetables with identity

Spanish vegetable cooking is not an afterthought. It may be smoky from a grill, silky from a long sofrito, crisp from frying, or brightened with vinegar. Match the method to the vegetable’s water content. Eggplant benefits from salting only when bitterness or excess moisture is a concern; modern varieties may need little treatment. Potatoes can be boiled, roasted, fried, or absorbed into a stew, and each method creates a different dish.

## Recipe: Pisto manchego

### Recipe information

- **Prep Time:** 25 minutes
- **Cook Time:** 45 minutes
- **Total Time:** About 1 hour 10 minutes
- **Servings:** 4–6
- **Difficulty:** Beginner
- **Cuisine:** Castilla-La Mancha-inspired

### Ingredients

- 60 ml (¼ cup) olive oil
- 1 onion, diced
- 1 red pepper and 1 green pepper, diced
- 1 small zucchini, diced
- 1 small eggplant, diced
- 500 g (1 lb 2 oz) crushed tomatoes
- 2 garlic cloves, minced
- Salt and black pepper
- Eggs, optional, to serve

### Instructions

1. Cook the onion and peppers in the oil over medium-low heat for 12–15 minutes.
2. Add the eggplant and zucchini. Cook until softened and lightly colored.
3. Add garlic and tomatoes. Simmer uncovered 20–25 minutes until thick and glossy.
4. Season and serve warm, at room temperature, or topped with a fried or poached egg.

### Chef’s tips

Cook watery vegetables in stages so the pan does not turn into a boil. Pisto improves after a short rest and can be refrigerated for 3–4 days.

## Recipe: Escalivada

Roast 2 red peppers, 1 eggplant, and 1 onion at 220°C (425°F) until charred outside and soft inside. Cover briefly to steam, peel the peppers and eggplant, and tear everything into strips. Dress with olive oil, sherry vinegar, and salt. The goal is smoky softness, not a uniform purée.

## Recipe: Spanish roasted potatoes

Toss 1 kg (2¼ lb) potato wedges with 60 ml olive oil, 4 crushed garlic cloves, rosemary, salt, and a little pimentón. Roast at 220°C (425°F) for 35–45 minutes, turning once, until crisp outside and tender inside. Do not coat the potatoes with wet tomato sauce before roasting; add a sauce at the table or after the crust develops.

## Other vegetable preparations

Garlic green beans should be blanched or steamed until just tender, then tossed in warm olive oil with sliced garlic and perhaps toasted almonds. Stuffed peppers require a filling that is already cooked or nearly cooked because the pepper’s skin may soften before a raw filling is safe. Eggplant with tomato can be fried, roasted, or braised; choose one texture and keep the sauce concentrated. Lentils with vegetables belong here as a main course or side, depending on portion and accompaniment.

# CHAPTER 26 — SPANISH BREADS AND SAVOURY PASTRIES

## Bread as an ingredient

Bread appears on the table, under a topping, inside a filling, and in sauces and desserts. Good bread needs flour, water, fermentation, heat, and time. A home oven will not behave like a stone hearth, but steam, a fully heated tray, and proper cooling can improve the crust.

## Recipe: Pan con tomate

### Recipe information

- **Prep Time:** 10 minutes
- **Cook Time:** 5 minutes
- **Total Time:** 15 minutes
- **Servings:** 4
- **Difficulty:** Beginner
- **Cuisine:** Catalan

### Ingredients

- 4 thick slices rustic bread
- 2 ripe tomatoes
- 1 garlic clove, optional
- Extra-virgin olive oil
- Flaky salt

### Instructions

1. Toast the bread until crisp at the edges but still substantial.
2. Rub lightly with garlic if using.
3. Cut the tomatoes in half and grate the flesh over a bowl, leaving the skins behind, or rub the cut side directly over the bread.
4. Spoon or spread the tomato over the toast. Finish with olive oil and salt just before serving.

### Cultural note

The dish is simple because the ingredients must be good. Watery, pale tomatoes and soft bread cannot be repaired with more oil.

## Recipe: Empanada gallega

### Recipe information

- **Prep Time:** 45 minutes plus resting
- **Cook Time:** 40–45 minutes
- **Total Time:** About 2 hours
- **Servings:** 8
- **Difficulty:** Intermediate
- **Cuisine:** Galician

### Ingredients

**Dough**

- 500 g (1 lb 2 oz) bread or plain flour
- 10 g (2 teaspoons) fine salt
- 7 g (2¼ teaspoons) instant yeast
- 250 ml (1 cup) warm water
- 60 ml (¼ cup) olive oil

**Filling**

- 45 ml (3 tablespoons) olive oil
- 2 onions, sliced
- 1 red pepper, diced
- 2 garlic cloves, minced
- 400 g (14 oz) crushed tomatoes
- 250 g (9 oz) drained tuna or cooked fish
- 1 teaspoon sweet pimentón

### Instructions

1. Mix flour, salt, and yeast. Add water and oil, knead until smooth, cover, and rest 60 minutes.
2. Cook the onions and pepper in oil until soft. Add garlic, tomato, and pimentón; reduce until thick. Cool, then fold in the fish.
3. Divide the dough. Roll one piece into a round or rectangle and place on a lined tray. Spread the cooled filling, leaving a border.
4. Cover with the second piece, seal, and cut a small steam vent. Brush with oil or beaten egg.
5. Bake at 200°C (400°F) for 40–45 minutes until deeply golden. Cool 15 minutes before slicing.

### Chef’s tips

The filling must be cool and fairly dry or it will make the dough soggy. A small vent prevents the top from ballooning. Empanada is useful for entertaining because it can be served warm or at room temperature.

## Empanadillas and olive bread

Empanadillas are smaller filled pastries that may be baked or fried. Seal the edges firmly and do not overfill. A regional olive bread may include olives, herbs, and oil in a lean dough; keep the olives well drained so they do not disrupt the crumb. Savory vegetable pastry can use the empanada filling principle with seasonal greens, roasted pepper, or artichoke.

# CHAPTER 27 — SPANISH DESSERTS

## Sweetness with restraint

Spanish desserts often use eggs, milk, almonds, citrus, cinnamon, olive oil, and bread. Their textures range from crisp fried dough to trembling custard. The best dessert chapters teach temperature control: eggs set gradually, sugar caramelizes quickly, and fried dough needs oil at a stable temperature.

## Recipe: Churros with chocolate

### Recipe information

- **Prep Time:** 15 minutes
- **Cook Time:** 20 minutes
- **Total Time:** 35 minutes
- **Servings:** 4
- **Difficulty:** Intermediate
- **Cuisine:** Spanish

### Ingredients

- 250 ml (1 cup) water
- 30 g (2 tablespoons) butter or olive oil
- 1 tablespoon sugar
- ½ teaspoon salt
- 150 g (5¼ oz) plain flour
- Frying oil
- Sugar and cinnamon, to finish

**Chocolate**

- 500 ml (2 cups) milk
- 120 g (4¼ oz) dark chocolate, chopped
- 1 tablespoon cornstarch mixed with 2 tablespoons cold water

### Instructions

1. Bring the water, butter, sugar, and salt to a simmer. Remove from the heat, add the flour at once, and stir until a smooth dough forms.
2. Return to low heat for 1 minute to dry slightly. Cool until safe to handle, then transfer to a sturdy piping bag fitted with a star nozzle.
3. Heat oil to 175°C (350°F). Pipe 10–12 cm (4–5 in) lengths into the oil, cutting with scissors. Fry in small batches until golden and crisp.
4. Drain, toss with cinnamon sugar, and serve.
5. For the chocolate, heat the milk and chocolate together. Whisk in the cornstarch slurry and cook until thick enough to coat a spoon.

### Safety note

Never pipe wet dough into oil with your hand near the surface, and never leave hot oil unattended. If the dough is too stiff to pipe, add a small amount of warm water only before frying and test one piece.

## Recipe: Crema catalana

Whisk 6 egg yolks, 120 g (4¼ oz) sugar, and 35 g (1¼ oz) cornstarch. Heat 750 ml (3 cups) milk with lemon peel and a cinnamon stick until steaming. Temper the milk into the yolks, return to the saucepan, and cook over medium-low heat while stirring until thick. Remove the aromatics, pour into shallow dishes, chill, and cover with a thin layer of sugar. Caramelize with a torch or hot grill immediately before serving.

## Recipe: Flan

### Recipe information

- **Prep Time:** 20 minutes
- **Cook Time:** 45–55 minutes
- **Total Time:** About 2 hours including cooling
- **Servings:** 6
- **Difficulty:** Intermediate
- **Cuisine:** Spanish

### Ingredients

- 150 g (5¼ oz) sugar for caramel
- 500 ml (2 cups) whole milk
- 4 large eggs
- 2 egg yolks
- 80 g (⅓ cup plus 1 tablespoon) sugar
- 1 teaspoon vanilla

### Instructions

1. Melt the caramel sugar in a saucepan until amber. Pour into a 1 litre (1 quart) dish or six ramekins and swirl carefully to coat.
2. Warm the milk with vanilla. Whisk eggs, yolks, and sugar without creating too much foam. Temper in the warm milk and strain.
3. Pour into the caramel-lined dish. Set it in a roasting pan and add hot water halfway up the sides.
4. Bake at 160°C (325°F) until the edges are set and the center trembles gently, 45–55 minutes. Cool in the water bath, then chill at least 4 hours before unmoulding.

### Common mistakes

A high oven creates bubbles and a rubbery texture. Caramel that is too pale tastes flat; caramel that is black tastes bitter. The custard continues to set while cooling, so do not wait for the center to become firm in the oven.

## Recipe: Tarta de Santiago

Mix 250 g (9 oz) finely ground almonds, 200 g (7 oz) sugar, 4 eggs, lemon zest, and a pinch of cinnamon. Spread in a buttered 22 cm (9 in) tin and bake at 175°C (350°F) for 30–35 minutes until golden and set. Cool completely and dust with icing sugar. The cake should be moist and almond-forward, not dry.

## Bread, rice, and citrus desserts

Torrijas transform day-old bread with milk, cinnamon, citrus, and egg before shallow frying or baking. Rice pudding depends on gentle stirring and gradual absorption; add citrus peel without the bitter white pith. Almond cake is forgiving when weighed accurately, while a citrus dessert should use acidity to keep sweetness lively. Regional custard pastries and seasonal fruit desserts belong to the same larger principle: keep the main flavor clear and use sugar to support it rather than overwhelm it.

## Closing note

Spanish cooking asks the cook to notice time. Let the sofrito become sweet. Let the beans become tender. Let the rice absorb without unnecessary stirring. Let a custard set slowly and a fried food drain properly. These are not delays between the important steps; they are the steps.

Across these chapters, the same lesson appears in different clothing. A cold soup and a hot stew both need balance. A tapa and a family pot both need generosity. A paella and an empanada both depend on managing moisture. Regional Spanish cooking becomes more understandable when the cook looks for the method beneath the name.

**PHASE COMPLETE — READY FOR THE NEXT PHASE**
"""


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_markdown_document(doc, SPANISH_CUISINE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()