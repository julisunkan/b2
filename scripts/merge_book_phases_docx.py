from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ASSETS = Path("attached_assets")
OUTPUT = ASSETS / "mastering_french_spanish_cooking_complete_manuscript.docx"
PHASE_FILES = [
    ASSETS / f"mastering_french_spanish_cooking_phase{number}_{suffix}.docx"
    for number, suffix in [
        (1, "blueprint"),
        (2, "front_matter"),
        (3, "foundations"),
        (4, "french_cuisine"),
        (5, "spanish_cuisine"),
        (6, "advanced_techniques"),
        (7, "menu_planning"),
        (8, "reference"),
        (9, "editorial_review"),
    ]
]


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(115, 115, 115)
    run._r.append(OxmlElement("w:fldChar"))
    run._r[-1].set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._r.append(instruction)
    run._r.append(OxmlElement("w:fldChar"))
    run._r[-1].set(qn("w:fldCharType"), "end")


def append_body(source, target):
    target_body = target.element.body
    target_section_properties = target_body.sectPr
    for child in source.element.body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        target_body.insert(target_body.index(target_section_properties), deepcopy(child))


def add_phase_break(target):
    paragraph = target.add_paragraph()
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()
    paragraph.paragraph_format.page_break_before = True


def configure_merged_header_and_footer(document):
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = (
            "MASTERING THE ART OF FRENCH & SPANISH COOKING  •  COMPLETE MANUSCRIPT"
        )
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(115, 115, 115)

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        add_page_number(footer_paragraph)


def main():
    missing = [path for path in PHASE_FILES if not path.exists()]
    if missing:
        names = ", ".join(path.name for path in missing)
        raise FileNotFoundError(f"Missing phase document(s): {names}")

    merged = Document(PHASE_FILES[0])
    for source_path in PHASE_FILES[1:]:
        add_phase_break(merged)
        append_body(Document(source_path), merged)

    configure_merged_header_and_footer(merged)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    merged.save(OUTPUT)
    print(
        f"Merged {len(PHASE_FILES)} phase documents into "
        f"{OUTPUT} ({OUTPUT.stat().st_size:,} bytes)"
    )


if __name__ == "__main__":
    main()