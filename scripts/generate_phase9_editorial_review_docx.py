from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from generate_book_blueprint_docx import configure_document, add_markdown_document


OUTPUT = Path(
    "attached_assets/mastering_french_spanish_cooking_phase9_editorial_review.docx"
)


def add_centered_text(doc, text, size=12, bold=False, italic=False, color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_title_page(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(105)
    run = paragraph.add_run("MASTERING THE ART OF\nFRENCH & SPANISH COOKING")
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(107, 62, 38)

    add_centered_text(
        doc,
        "A Regional, Technique-Driven Guide to the Ingredients, "
        "Traditions, and Recipes of France and Spain",
        size=14,
        italic=True,
        color="8A5A3B",
    )
    add_centered_text(doc, "PHASE 9 — FINAL EDITORIAL REVIEW", size=12, bold=True)
    add_centered_text(doc, "By [AUTHOR NAME]", size=12)
    doc.add_page_break()


EDITORIAL_REVIEW = r"""
# PHASE 9 — FINAL EDITORIAL REVIEW

## Purpose and editorial status

This phase is the final editorial audit of the assembled manuscript for **MASTERING THE ART OF FRENCH & SPANISH COOKING**. It is not a replacement for line editing, recipe testing, copyediting, proofreading, or legal review. It is the control document that identifies what must be corrected before those final production stages can be trusted.

The audit compares the current phase files against the approved book blueprint and the master writing prompt. It checks the manuscript as a reader will encounter it: as one book with a continuous table of contents, consistent chapter numbering, repeatable recipe formatting, reliable culinary language, and a clear progression from foundations to reference material.

### Editorial verdict

**Status: REVISION REQUIRED BEFORE PUBLICATION.**

The manuscript has a strong educational concept, a useful regional scope, and substantial material across the first eight phases. It is not yet publication-ready because the current assembled files contain structural numbering conflicts, an incomplete recipe count relative to the blueprint, and unresolved consistency questions that should be settled before a final contents page, index, or print layout is produced.

This verdict is constructive rather than final in tone. The core manuscript can be made coherent by resolving the architecture first, then running the recipe, cultural, language, and production checks in that order.

## 1. Audit basis and source inventory

The review was performed against:

- The original master prompt supplied with the project
- The Phase 1 book blueprint
- The Phase 2 front matter
- The Phase 3 kitchen foundations
- The Phase 4 French cuisine section
- The Phase 5 Spanish cuisine section
- The Phase 6 advanced techniques
- The Phase 7 menu-planning section
- The Phase 8 reference section

The audit treats the DOCX files as phase deliverables that still need to be merged into one controlled manuscript. A heading scan is useful for finding structural risks, but it does not replace a human culinary test of every recipe.

### Current visible inventory

- French cuisine is presented in Chapters 1–12.
- Spanish cuisine is presented in Chapters 18–27.
- Advanced techniques are presented in Chapters 28–36.
- Menu planning is presented in Chapters 36–42.
- The reference section is presented in Chapters 39–46.
- The phase files visibly contain 18 French recipe headings and 24 Spanish recipe headings, plus 10 foundational preparations in the foundations plan. These counts must be reconciled against the blueprint target of approximately 138 total recipe entries.
- Phase 8 includes both a recipe index and an ingredient index, but both indexes must be regenerated after chapter renumbering and recipe additions or removals.

## 2. Structural audit

### Findings

#### Finding S1 — Missing chapter range

The French section ends at Chapter 12 and the Spanish section begins at Chapter 18. Chapters 13–17 are therefore absent from the assembled sequence. The blueprint and master prompt identify Spanish cuisine as Part II beginning with the Spanish understanding chapter; the final manuscript must choose one authoritative numbering plan and apply it everywhere.

#### Finding S2 — Duplicate Chapter 36

Phase 6 uses Chapter 36 for “Technique Integration and Diagnostic Cooking,” while Phase 7 uses Chapter 36 for “A Classic French Dinner.” One of these headings must be renumbered after the master sequence is approved.

#### Finding S3 — Menu/reference collision

Phase 7 occupies Chapters 36–42, while Phase 8 begins at Chapter 39. This creates duplicate chapter numbers and makes the table of contents, cross-references, and indexes unreliable.

#### Finding S4 — Repeated phase labels

The Phase 7 file contains a repeated “PHASE 7 — MENU PLANNING” heading. Repeated labels can be retained only when one is clearly a cover-page element and the other is a section heading; otherwise remove the duplicate.

#### Finding S5 — Contents cannot yet be final

No table of contents, page references, internal chapter references, or index entries should be frozen until the corrected master sequence is merged. Page numbers generated before the structural pass will become stale.

### Required structural correction order

1. Approve one master chapter map.
2. Assign unique chapter numbers to all chapters from the foundations through the reference section.
3. Update every chapter heading, contents entry, menu cross-reference, recipe index entry, and ingredient-index entry.
4. Merge the phase documents into a single working manuscript.
5. Regenerate the table of contents and both indexes from the merged manuscript.
6. Perform a second heading scan for duplicate numbers, missing numbers, and inconsistent title text.

### Recommended master sequence

The cleanest repair is to preserve the chapter titles and correct the numbering consecutively:

- Foundations: Chapters 1–6
- French cuisine: Chapters 7–18
- Spanish cuisine: Chapters 19–28
- Advanced techniques: Chapters 29–37
- Menu planning: Chapters 38–44
- Reference: Chapters 45–52

This sequence follows the detailed blueprint, avoids collisions, and leaves no unexplained gaps. If the editorial team prefers to retain the phase files’ existing chapter numbers, it must still create a new unique sequence rather than carrying the current collisions forward.

## 3. Completeness audit

### Coverage against the brief

The manuscript addresses the major promised areas: equipment, knife skills, cooking techniques, French and Spanish regional cooking, advanced technique lessons, menu planning, ingredients, substitutions, dictionary terms, conversions, seasonality, and a beginner learning path.

The following items require explicit completion checks:

- Every chapter named in the approved blueprint is present after renumbering.
- Every chapter has an introduction, a clear teaching purpose, and a closing transition where appropriate.
- Every recipe named as a required example is either included or deliberately replaced with an editorially documented equivalent.
- The final recipe count is reconciled by cuisine, chapter, difficulty, and recipe type.
- The master recipe index includes every recipe exactly once.
- The master ingredient index includes ingredients used in the final recipes, not merely ingredients discussed in reference prose.
- The final manuscript contains the copyright page and a cooking disclaimer.
- Photography suggestions are either retained as production markers or removed in favor of a separate art-direction document.

### Recipe-count reconciliation

The blueprint proposes approximately 128 primary cuisine recipes plus 10 foundational preparations, for approximately 138 recipe entries. The currently visible Phase 4 and Phase 5 headings account for 42 cuisine recipes, with the foundations plan identifying 10 preparations. This is a material gap, not a formatting detail.

Before publication, create a recipe ledger with these fields:

- Recipe title and original-language title where appropriate
- Country and region
- Phase and chapter
- Recipe type
- Difficulty
- Servings
- Presence of ingredients, equipment, instructions, tips, mistakes, variations, substitutions, serving, storage, and cultural note
- Index status
- Test status

Do not solve the gap by adding filler recipes. Add only recipes that teach a distinct technique, represent a meaningful regional preparation, or improve the balance of the book.

## 4. Recipe audit protocol

Every recipe should pass the following checks before it is marked final.

### Recipe structure

Confirm that each entry contains, in the same order:

1. Recipe title
2. Concise introduction
3. Recipe information
4. Equipment
5. Ingredients with plausible metric and U.S. quantities
6. Numbered instructions
7. Chef’s tips
8. Common mistakes
9. Variations
10. Substitutions where useful
11. Serving suggestions
12. Storage and reheating guidance
13. Cultural or regional note where appropriate

Optional sections should be omitted intentionally, not accidentally. A short recipe may still be complete; completeness is not the same as length.

### Ingredient-to-instruction reconciliation

For each recipe:

- Mark every listed ingredient at the step where it is used.
- Confirm that no instruction requires an ingredient that is absent from the list.
- Check that divided quantities are labeled in both ingredients and instructions.
- Confirm that preparation verbs match the ingredient state: sliced, diced, minced, drained, dried, thawed, or at room temperature.
- Check that salt contributed by stock, cured foods, cheese, olives, anchovies, or preserved ingredients is considered before final seasoning.

### Time, temperature, and equipment

- Check that preparation and cooking times are realistic in a normal home kitchen.
- Confirm that oven temperatures include both Celsius and Fahrenheit where relevant.
- Confirm that preheating is stated.
- Confirm that pan size and capacity are adequate for the quantity.
- Check that high-heat steps allow for smoke, splatter, and ventilation.
- Use sensory cues such as color, aroma, texture, bubbling, and resistance alongside time.
- Add a thermometer cue where food safety depends on internal temperature.

### Safety

Give special review to poultry, minced meat, shellfish, raw or lightly cooked eggs, deep-frying, hot sugar, pressure or high-heat cooking, and leftovers. Storage advice must distinguish between a dish that keeps well and one that loses quality quickly. Local food-safety guidance takes priority over a general cookbook instruction.

## 5. Culinary and cultural audit

### Accuracy principles

- Name the country and region when a dish has a strong regional identity.
- Distinguish a traditional preparation from a modern variation or home adaptation.
- Avoid calling one version universally authentic when established regional versions differ.
- Do not attribute a dish to a precise historical person, date, or event without reliable support.
- Explain French and Spanish terms when first introduced, then use them consistently.
- Do not flatten either cuisine into a small set of tourist clichés.

### Targeted cultural checks

#### French section

Confirm that regional discussion distinguishes, where relevant, the Atlantic coast, Normandy and Brittany, Alsace, Burgundy, Provence, Occitanie, the Loire, Bordeaux, the Alps, and Basque-influenced areas without presenting administrative borders as absolute culinary boundaries.

Review sauce terminology, stock terminology, pastry terminology, and the use of butter, cream, wine, mustard, herbs, and olive oil for both technical accuracy and regional context.

#### Spanish section

Confirm that the discussion distinguishes, where relevant, Andalusia, Catalonia, Valencia, Galicia, the Basque Country, Asturias, Madrid, Castilla y León, Castilla-La Mancha, Murcia, Aragón, Extremadura, and the Canary Islands.

Review the treatment of sofrito, pimentón, saffron, rice, olive oil, cured meats, tapas, raciones, and socarrat. Paella terminology deserves special care: not every Spanish rice dish is paella, and a recipe should identify its regional style rather than imply a single definitive national form.

### Cultural red flags to remove

- Unsupported origin stories presented as fact
- Statements that imply one cuisine is more authentic or refined than the other
- “Traditional” labels applied to newly invented substitutions
- Regional names used only as decorative adjectives
- Inconsistent transliteration, accenting, or capitalization of French and Spanish terms

## 6. Language and house-style audit

Choose the house style before the final copyedit and apply it globally.

### Recommended decisions

- Use metric first, followed by a practical U.S. equivalent where useful.
- Use one spelling policy for words such as “savory/savoury,” “color/colour,” and “center/centre.”
- Use one capitalization policy for dish names, headings, and ingredient names.
- Use one policy for accent marks, apostrophes, hyphens, and italicized foreign terms.
- Use “chef’s tips” consistently rather than alternating with “chef’s notes” unless the distinction is intentional.
- Use “storage” and “reheating” labels consistently.
- Standardize “extra-virgin olive oil,” “crème fraîche,” “pimentón,” “aioli/alioli,” and other recurring terms.

### Copyediting pass

Search the merged manuscript for:

- Double spaces and repeated punctuation
- Inconsistent curly and straight apostrophes
- Heading levels that skip or restart unexpectedly
- Recipe titles that vary between English and original-language forms
- Measurements that mix abbreviations and full words without a rule
- Duplicate sentences or repeated introductory paragraphs
- Place names and diacritics used inconsistently
- “Beginner,” “intermediate,” and “advanced” labels that do not match recipe complexity

## 7. Publishing and production audit

Before export:

1. Merge the approved text into one source manuscript.
2. Apply one style system to body text, headings, recipe metadata, lists, notes, and photo suggestions.
3. Insert an automatically generated table of contents.
4. Update fields and verify that page numbers correspond to the final layout.
5. Regenerate the recipe and ingredient indexes from final text.
6. Check widows, orphans, blank pages, orphaned headings, broken lists, and recipe continuations.
7. Confirm that photography markers are either placed consistently or removed from reader-facing copy.
8. Check the title page, copyright page, disclaimer, dedication, and author placeholder.
9. Export a proof PDF and inspect it at print size and on a screen.
10. Preserve an editable DOCX source and a final proof with a clear version name.

### Production acceptance criteria

The manuscript is ready for production only when:

- Every chapter has one unique number.
- The contents page matches the body.
- Every index entry points to the correct final page.
- The recipe ledger matches the manuscript.
- No recipe has unresolved ingredient or instruction contradictions.
- All required safety and storage notes are present.
- The language and measurement house style is consistent.
- Cultural claims have been checked or cautiously phrased.
- A human editor and a competent cook have both reviewed the proof.

## 8. Prioritized revision plan

### P0 — Release blockers

- Resolve the chapter-numbering collisions and missing range.
- Merge the phases into one controlled manuscript.
- Reconcile the recipe count and required chapter coverage.
- Rebuild the contents page and indexes after structural changes.
- Resolve any recipe safety or ingredient/instruction contradiction found during testing.

### P1 — High-impact editorial corrections

- Run the complete recipe ledger and format audit.
- Standardize measurements, terminology, headings, and spelling.
- Review cultural and historical claims, especially regional labels and paella terminology.
- Check transitions between cuisine sections, techniques, menus, and reference material.
- Confirm that the beginner learning path points to recipes that actually exist.

### P2 — Finishing and production polish

- Perform a line edit for rhythm, repetition, and unnecessary filler.
- Standardize photography suggestions and art-direction language.
- Complete the print and digital proof checks.
- Verify accessibility of headings, lists, contrast, and navigational structure in the digital edition.

## 9. Final editorial sign-off sheet

Complete this sign-off only after the P0 and P1 findings are closed.

- [ ] Master chapter map approved
- [ ] All phase files merged into one manuscript
- [ ] No missing or duplicate chapters
- [ ] Recipe ledger complete
- [ ] Recipe count approved by the editor
- [ ] Every recipe tested or explicitly marked for testing
- [ ] Ingredient and instruction reconciliation complete
- [ ] Safety and storage review complete
- [ ] Cultural review complete
- [ ] House style sheet applied
- [ ] Table of contents regenerated
- [ ] Recipe index regenerated
- [ ] Ingredient index regenerated
- [ ] Copyright and disclaimer approved
- [ ] Photography markers resolved
- [ ] Print proof reviewed
- [ ] Digital proof reviewed
- [ ] Final editorial approval recorded

## Closing assessment

This manuscript has the ingredients of a useful, distinctive cookbook: it teaches technique, respects regional identity, and aims to connect recipes with the decisions that make them work. The next editorial effort should resist cosmetic polishing until the structure and recipe ledger are settled. Once the architecture is coherent, the remaining work can be measured, assigned, tested, and signed off with confidence.

**PHASE COMPLETE — READY FOR THE NEXT PHASE**
"""


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_markdown_document(doc, EDITORIAL_REVIEW)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()