from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("attached_assets/mastering_french_spanish_cooking_phase1_blueprint.docx")
SOURCE_PROMPT = Path(
    "attached_assets/"
    "Pasted--MASTER-PROMPT-FOR-WRITING-THE-BOOK-BOOK-TITLE-MASTERIN_1786779460863.txt"
)


BLUEPRINT = r"""
# PHASE 1 — BOOK BLUEPRINT

## 1. Final Title

# MASTERING THE ART OF FRENCH & SPANISH COOKING

### Subtitle

**A Regional, Technique-Driven Guide to the Ingredients, Traditions, and Recipes of France and Spain**

### Working author line

**By [AUTHOR NAME]**

The author name remains intentionally blank until one is provided.

## 2. Target Audience

This cookbook is designed for:

- Beginners who want to learn foundational cooking skills
- Intermediate home cooks seeking greater confidence
- Experienced cooks interested in authentic European traditions
- Readers who want to understand the reasoning behind classic techniques
- Food enthusiasts interested in French and Spanish regional cuisine
- International English-speaking readers who need both metric and U.S. measurements
- Home entertainers who want to plan complete French and Spanish meals
- Readers interested in ingredients, food history, culinary terminology, and regional identity

The book should remain accessible without becoming simplistic. Beginners should be able to follow the recipes, while experienced cooks should find useful technical explanations and regional context.

## 3. Book Promise

By the end of this book, readers will be able to:

1. Understand the central principles of French and Spanish cooking.
2. Recognize the major regional identities within both countries.
3. Select, prepare, and substitute important ingredients intelligently.
4. Use essential kitchen equipment safely and effectively.
5. Execute fundamental techniques such as sautéing, searing, braising, emulsifying, reducing, frying, roasting, and baking.
6. Prepare classic French and Spanish dishes with accurate methods and realistic home-kitchen instructions.
7. Understand why a technique works rather than merely following directions.
8. Build sauces, stocks, flavor bases, menus, and complete meals.
9. Distinguish traditional preparations from modern variations.
10. Progress from simple dishes to advanced recipes with confidence.

The book should feel like a complete culinary education rather than a collection of disconnected recipes.

## 4. Scope and Editorial Targets

### Manuscript length

- **Target:** approximately 100,000–105,000 words
- **Acceptable range:** 80,000–110,000 words
- **Estimated print length:** approximately 280–330 pages, depending on trim size, photography, typography, and recipe layout

### Recipe count

- **128 primary cuisine recipes**
  - 64 French recipes
  - 64 Spanish recipes
- **10 foundational preparations**
  - French stocks, sauces, and bases
  - Spanish sofrito and related foundations
- **Total:** approximately **138 recipe entries**

### Recipe balance

The book will maintain an approximately equal French-Spanish split while allowing each cuisine to retain its own structure and character.

### Measurement policy

Recipes will use:

- Metric measurements first where practical
- Common U.S. equivalents in parentheses
- Celsius and Fahrenheit oven temperatures
- Weight-based measurements for baking and ingredients where precision matters
- Volume measurements where home cooks commonly rely on them

# 5. Detailed Table of Contents

## FRONT MATTER

### Title Page

- Book title
- Subtitle
- Author line

### Copyright and Disclaimer

- Standard copyright notice
- Publication details placeholder
- Recipe and food-safety disclaimer
- Statement that cooking times and temperatures may vary by equipment and ingredients
- Allergen and personal-health disclaimer

### Dedication

A brief, elegant dedication celebrating cooks, teachers, family meals, and the pleasure of sharing food.

### Table of Contents

A complete hierarchical contents page covering all parts, chapters, major sections, and reference material.

## INTRODUCTION

### The Shared Table of France and Spain

- What French cooking represents
- What Spanish cooking represents
- Why both cuisines reward careful study
- Shared Mediterranean, Atlantic, agricultural, and regional influences
- Important distinctions in ingredients, methods, sauces, seasoning, and dining culture
- Why regional identity matters
- How technique transforms ordinary ingredients
- How to use this book progressively
- Encouragement for the reader

## HOW TO USE THIS BOOK

### Reading the Chapters

- The educational sequence
- How the French and Spanish sections are organized
- How the technique sections support the recipes

### Understanding Recipe Information

- Prep time
- Cook time
- Total time
- Servings
- Difficulty
- Cuisine and regional notes

### Measurements and Temperatures

- Metric and U.S. measurements
- Weight versus volume
- Oven temperature interpretation
- Variations in equipment and cookware

### Mise en Place

- Preparing ingredients before cooking
- Organizing equipment
- Reading the entire recipe before beginning
- Separating preparation from execution

### Adapting Recipes Responsibly

- Sensible substitutions
- What may be changed safely
- What defines the character of a dish
- How to preserve balance when adapting

# KITCHEN FOUNDATIONS

## Chapter 1 — Essential Kitchen Equipment

### Core Equipment

- Chef’s knife
- Paring knife
- Bread knife
- Cutting boards
- Saucepan
- Frying pan
- Sauté pan
- Stockpot
- Dutch oven
- Roasting pan
- Baking trays
- Mixing bowls

### Essential Hand Tools

- Whisk
- Wooden spoon
- Spatula
- Tongs
- Ladle
- Fine-mesh strainer
- Colander
- Grater
- Peeler
- Kitchen scale
- Measuring cups and spoons
- Instant-read thermometer

### Optional Equipment

- Mandoline
- Blender
- Food processor
- Stand mixer
- Pastry tools
- Paella pan
- Mortar and pestle
- Grill pan

### Choosing, Maintaining, and Storing Equipment

- Materials and durability
- Heat distribution
- Cleaning
- Storage
- Equipment substitutions for small kitchens

**Estimated length:** 2,000 words  
**Recipes:** None

## Chapter 2 — Knife Skills and Kitchen Safety

### Knife Control

- Proper grip
- The guiding hand
- Stable cutting surfaces
- Safe working posture

### Fundamental Cuts

- Chopping
- Dicing
- Slicing
- Julienne
- Brunoise
- Chiffonade
- Mincing
- Carving

### Knife Maintenance

- Honing
- Sharpening
- Safe storage
- Recognizing a dull blade

### General Kitchen Safety

- Preventing cuts
- Managing heat and steam
- Avoiding cross-contamination
- Safe handling of raw meat, seafood, and eggs
- Handling hot oil
- Using a thermometer

**Estimated length:** 1,500 words  
**Recipes:** None

## Chapter 3 — Essential Cooking Techniques

Each technique will explain:

1. What it is
2. Why it works
3. How to perform it
4. Common mistakes
5. Sensory cues for success

### Dry-Heat Techniques

- Sautéing
- Sweating
- Searing
- Roasting
- Baking
- Grilling
- Broiling
- Pan-frying
- Deep-frying

### Moist-Heat Techniques

- Boiling
- Simmering
- Poaching
- Blanching
- Stewing
- Braising

### Sauce and Flavor Techniques

- Deglazing
- Reducing
- Emulsifying
- Caramelizing
- Confit
- Flambéing
- Slow cooking

**Estimated length:** 4,000 words  
**Recipes:** None

## Chapter 4 — Flavor Fundamentals

### The Main Components of Flavor

- Salt
- Acidity
- Fat
- Sweetness
- Bitterness
- Umami
- Herbs
- Spices
- Aromatics

### Browning and the Maillard Reaction

- Why browned food tastes deeper
- Proper pan temperature
- Moisture management
- Avoiding overcrowding

### Balance and Adjustment

- Correcting excessive saltiness
- Correcting acidity
- Adding richness
- Brightening heavy dishes
- Building layered flavor

**Estimated length:** 2,500 words  
**Recipes:** None

## Chapter 5 — Stocks, Broths, and Foundational Preparations

### French Foundations

1. Basic vegetable stock
2. Chicken stock
3. Beef stock
4. Fish stock
5. Court-bouillon
6. Mirepoix
7. Bouquet garni

### Spanish Foundations

8. Basic sofrito
9. Spanish seafood broth
10. Picada

### Using Foundations in Later Recipes

- Storage
- Freezing
- Degreasing
- Clarifying
- Adjusting strength
- Choosing the appropriate foundation

**Estimated length:** 2,500 words  
**Recipes:** 10 foundational preparations

## Chapter 6 — Ingredient Fundamentals

### Buying and Storing Ingredients

- Seasonality
- Freshness
- Pantry organization
- Refrigeration
- Freezing
- Reading labels

### Core Ingredient Categories

- Fats
- Aromatics
- Herbs
- Spices
- Grains
- Legumes
- Meat
- Poultry
- Seafood
- Dairy
- Eggs
- Flour and sugar

### Substitution Principles

- Flavor
- Texture
- Moisture
- Cooking time
- Regional authenticity

**Estimated length:** 2,500 words  
**Recipes:** None

# PART I — THE ART OF FRENCH COOKING

## Chapter 7 — Understanding French Cuisine

### Culinary Philosophy

- Precision and restraint
- Technique as a path to consistency
- The importance of sauces and foundations
- Seasonality
- Balance and presentation

### Historical Development

- Classical culinary traditions
- The evolution of restaurant cooking
- Home-style cuisine
- Modern French cooking
- Influence from neighboring cultures

### Regional France

- Île-de-France
- Normandy
- Brittany
- Alsace
- Burgundy
- Provence
- Occitanie
- Loire Valley
- Bordeaux
- The French Alps
- Basque-influenced areas

Regional boundaries will be presented as overlapping culinary traditions rather than rigid administrative categories.

**Estimated length:** 2,000 words  
**Recipes:** None

## Chapter 8 — Essential French Ingredients

### Fats and Dairy

- Butter
- Olive oil
- Cream
- Cheese
- Eggs

### Aromatics and Vegetables

- Garlic
- Shallots
- Onions
- Leeks
- Carrots
- Celery
- Mushrooms
- Tomatoes
- Potatoes

### Herbs and Seasonings

- Tarragon
- Thyme
- Rosemary
- Parsley
- Chervil
- Bay leaf
- Dijon mustard
- Vinegar
- Wine

### Proteins and Legumes

- Lentils
- Beans
- Seafood
- Poultry
- Beef
- Pork
- Lamb

Each ingredient entry will explain selection, storage, culinary purpose, traditional uses, and practical substitutions.

**Estimated length:** 2,000 words  
**Recipes:** None

## Chapter 9 — The Foundations of French Sauce

### Classic Sauce Families

- Béchamel
- Velouté
- Espagnole
- Tomato sauce
- Hollandaise

### Sauce Techniques

- Making a roux
- Deglazing
- Reducing
- Mounting with butter
- Building a pan sauce
- Creating cream sauces
- Making mustard sauces
- Herb sauces
- Vinaigrettes

### Troubleshooting

- Lumps
- Splitting
- Excessive thickness
- Thin sauces
- Over-reduction
- Under-seasoning

**Estimated length:** 2,500 words  
**Recipes:** Core sauce formulas integrated into the chapter

## Chapter 10 — French Breakfast and Brunch

Planned recipes:

1. Croissants
2. Pain perdu
3. Basic crêpes
4. Buckwheat galettes
5. French omelet
6. Brioche
7. Savory baked eggs

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 11 — French Appetizers and Hors d’Oeuvres

Planned recipes:

1. Gougères
2. Œufs mayonnaise
3. Country-style pâté
4. Vegetable terrine
5. French savory tart
6. Cheese and herb tartlets
7. Provençal stuffed vegetables

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 12 — French Soups and Salads

Planned recipes:

1. French onion soup
2. Vichyssoise
3. Seasonal potage
4. Soupe au pistou
5. Salade Niçoise
6. French lentil salad
7. Warm vegetable salad

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 13 — French Poultry and Meat

### Regional Context

- Burgundy and wine-based cooking
- Normandy and cream
- Provence and aromatic herbs
- Southwest France and confit traditions

Planned recipes:

1. Coq au vin
2. Coq au vin blanc
3. Chicken fricassée
4. Duck confit
5. Duck breast with pan sauce
6. Beef bourguignon
7. Steak au poivre
8. Blanquette de veau
9. Provençal lamb stew
10. French braised pork

**Estimated length:** 4,500 words  
**Recipes:** 10

## Chapter 14 — French Seafood

Planned recipes:

1. Moules marinières
2. Bouillabaisse-style fish stew
3. Fish meunière
4. Provençal baked cod
5. Salmon with sorrel-style sauce
6. Garlic shrimp
7. Seared scallops with beurre blanc

Regional distinctions between Atlantic, Mediterranean, and coastal home cooking will be explained.

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 15 — French Vegetables and Side Dishes

Planned recipes:

1. Ratatouille
2. Gratin dauphinois
3. Pommes purée
4. Pommes boulangère
5. Haricots verts with shallots
6. Glazed carrots
7. French-style mushrooms
8. Braised leeks

**Estimated length:** 3,500 words  
**Recipes:** 8

## Chapter 16 — French Breads and Pastries

### Dough Fundamentals

- Yeast
- Gluten
- Fermentation
- Lamination
- Proofing
- Pastry temperature control

Planned recipes:

1. Baguette
2. Brioche
3. Croissants
4. Pain au chocolat
5. Fougasse
6. Pâte brisée
7. Shortcrust pastry
8. Choux pastry

**Estimated length:** 4,000 words  
**Recipes:** 8

## Chapter 17 — French Desserts

Planned recipes:

1. Crème brûlée
2. Crème caramel
3. Mousse au chocolat
4. Tarte Tatin
5. Clafoutis
6. Madeleines
7. Profiteroles
8. Chocolate tart
9. Seasonal fruit tart
10. Île flottante

**Estimated length:** 4,500 words  
**Recipes:** 10

### French recipe total: 64 recipes

# PART II — THE ART OF SPANISH COOKING

## Chapter 18 — Understanding Spanish Cuisine

### Culinary History and Influences

- Mediterranean traditions
- Atlantic traditions
- Moorish influence
- Agricultural practices
- Seafood culture
- Olive oil
- Communal dining
- Tapas culture

### Regional Spain

- Andalusia
- Catalonia
- Valencia
- Galicia
- Basque Country
- Asturias
- Madrid
- Castilla y León
- Castilla-La Mancha
- Murcia
- Aragón
- Extremadura
- Canary Islands

Spanish cuisine will be treated as a collection of highly regional traditions rather than a single uniform style.

**Estimated length:** 2,000 words  
**Recipes:** None

## Chapter 19 — Essential Spanish Ingredients

### Foundational Ingredients

- Extra-virgin olive oil
- Garlic
- Onions
- Tomatoes
- Peppers
- Smoked paprika
- Saffron
- Sherry vinegar
- Olives
- Almonds
- Potatoes
- Rice

### Legumes and Preserved Foods

- Beans
- Chickpeas
- Lentils
- Cured meats
- Chorizo
- Spanish cheeses

### Seafood and Herbs

- Regional seafood
- Fresh herbs
- Preserved fish
- Salted and dried ingredients

Each entry will include taste, purpose, storage, availability, and realistic substitutions.

**Estimated length:** 2,000 words  
**Recipes:** None

## Chapter 20 — Tapas: The Heart of Spanish Dining

### Understanding Tapas

- Cultural role
- Small plates and shared dining
- Tapas, pinchos, and raciones
- Hot and cold preparations
- Building a balanced spread

Planned recipes:

1. Patatas bravas
2. Gambas al ajillo
3. Tortilla española
4. Potato and ham croquetas
5. Pimientos de Padrón
6. Albóndigas
7. Marinated olives
8. Garlic mushrooms
9. Spanish grilled vegetables

**Estimated length:** 4,000 words  
**Recipes:** 9

## Chapter 21 — Spanish Soups and Salads

Planned recipes:

1. Gazpacho
2. Salmorejo
3. Ajo blanco
4. Spanish bean soup
5. Lentil soup
6. Tomato and onion salad
7. Orange and olive salad

Regional differences in climate, ingredients, and serving temperature will be addressed.

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 22 — Paella and Spanish Rice Dishes

### Rice Fundamentals

- Choosing rice
- Stock preparation
- Saffron
- Pan shape
- Heat management
- Absorption
- Resting
- Socarrat
- When to stir and when not to stir

Planned recipes:

1. Paella Valenciana
2. Seafood paella
3. Mixed paella
4. Vegetable paella
5. Arroz negro
6. Catalan-style rice dish
7. Simple regional baked rice
8. Rice with beans and seasonal vegetables

The chapter will distinguish paella from other Spanish rice dishes and explain regional variation without presenting one version as universally definitive.

**Estimated length:** 4,000 words  
**Recipes:** 8

## Chapter 23 — Spanish Seafood

Planned recipes:

1. Garlic prawns
2. Galician-style octopus
3. Squid with onions
4. Steamed mussels
5. Clams in a sherry-style sauce
6. Salt-cod preparation
7. Sardines with herbs

Atlantic and Mediterranean seafood traditions will be compared.

**Estimated length:** 3,500 words  
**Recipes:** 7

## Chapter 24 — Spanish Meat and Poultry

Planned recipes:

1. Pollo al ajillo
2. Spanish meatballs
3. Catalan-style chicken
4. Pork with paprika
5. Spanish pork stew
6. Regional lamb stew
7. Beef braise
8. Slow-cooked chickpeas with meat
9. Madrid-style stew

**Estimated length:** 4,500 words  
**Recipes:** 9

## Chapter 25 — Spanish Vegetables and Side Dishes

Planned recipes:

1. Patatas bravas
2. Spanish roasted potatoes
3. Pisto
4. Escalivada
5. Garlic green beans
6. Stuffed peppers
7. Eggplant with tomato
8. Lentils with vegetables

**Estimated length:** 3,500 words  
**Recipes:** 8

## Chapter 26 — Spanish Breads and Savory Pastries

Planned recipes:

1. Pan con tomate
2. Rustic Spanish bread
3. Empanada
4. Empanadillas
5. Regional olive bread
6. Savory vegetable pastry

**Estimated length:** 3,000 words  
**Recipes:** 6

## Chapter 27 — Spanish Desserts

Planned recipes:

1. Churros
2. Crema catalana
3. Flan
4. Tarta de Santiago
5. Torrijas
6. Spanish rice pudding
7. Almond cake
8. Citrus dessert
9. Regional custard pastry
10. Seasonal fruit dessert

**Estimated length:** 4,500 words  
**Recipes:** 10

### Spanish recipe total: 64 recipes

# PART III — MASTERING THE TECHNIQUES

## Chapter 28 — Heat Control

- Low, medium, and high heat
- Pan temperature
- Oven temperature
- Carryover cooking
- Moist heat and dry heat
- Recognizing temperature through visual and sensory cues

## Chapter 29 — Building Flavor

- Browning
- Caramelization
- Maillard reaction
- Salt
- Acidity
- Fat
- Sweetness
- Bitterness
- Umami
- Herbs and spices

## Chapter 30 — Deglazing and Reduction

- Turning browned bits into sauce
- Choosing wine, stock, or vinegar
- Managing salt concentration
- Managing acidity
- Controlling texture

## Chapter 31 — Emulsions

- Mayonnaise
- Aioli
- Hollandaise
- Vinaigrettes
- Butter-based emulsions
- Diagnosing broken sauces

## Chapter 32 — Working with Seafood

- Selecting fresh seafood
- Safe storage
- Cleaning
- Deboning
- Preventing overcooking
- Doneness cues
- Cross-contamination prevention

## Chapter 33 — Working with Meat

- Selecting cuts
- Seasoning
- Searing
- Braising
- Resting
- Carving
- Doneness
- Carryover cooking

## Chapter 34 — Cooking Rice

- Rice-to-liquid ratios
- Short-, medium-, and long-grain rice
- Absorption
- Stirring versus not stirring
- Resting
- Texture control
- Troubleshooting undercooked or wet rice

## Chapter 35 — Baking Fundamentals

- Measuring flour
- Gluten
- Yeast
- Fermentation
- Eggs
- Butter
- Sugar
- Temperature
- Oven behavior
- Troubleshooting dough and pastry

**Estimated Part III length:** 3,500 words  
**New recipes:** None; techniques refer back to the main recipe chapters.

# PART IV — FRENCH AND SPANISH MENU PLANNING

## Chapter 36 — A Classic French Dinner

- Aperitif
- Starter
- Main course
- Side dish
- Dessert
- Beverage pairing
- Preparation schedule
- Serving order

## Chapter 37 — A Spanish Tapas Evening

- Number of dishes
- Hot and cold balance
- Vegetarian and seafood balance
- Serving sequence
- Beverage pairing
- Preparation schedule

## Chapter 38 — French Dinner for Two

- Elegant but manageable menu
- Shared preparation
- Timing for a small kitchen
- Plating and serving

## Chapter 39 — Spanish Family Feast

- Generous family-style service
- Make-ahead dishes
- Communal presentation
- Managing multiple components

## Chapter 40 — French Brunch

- Pastry
- Eggs
- Savory dishes
- Fruit
- Coffee and beverage service

## Chapter 41 — Mediterranean Seafood Dinner

- Selecting complementary seafood
- Building a coherent meal
- Avoiding repetition
- Timing delicate ingredients

## Chapter 42 — French and Spanish Celebration Menu

- Combining traditions respectfully
- Creating contrast without confusion
- Coordinating multiple courses
- Final service plan

**Estimated length:** 2,500 words  
**New recipes:** None; menus use recipes from previous chapters.

# PART V — INGREDIENTS AND SUBSTITUTIONS

## Chapter 43 — Ingredient Reference

For each important ingredient:

- Original name
- Flavor profile
- Culinary purpose
- Common uses
- Storage
- Availability
- Possible substitute
- Effect of substitution

## Chapter 44 — Responsible Substitution

- Dairy substitutions
- Oil substitutions
- Herb substitutions
- Seafood substitutions
- Meat substitutions
- Flour and grain substitutions
- Spice substitutions
- Ingredients that should not be casually replaced

**Estimated length:** 4,000 words

# PART VI — FRENCH AND SPANISH CULINARY DICTIONARY

## Chapter 45 — French Culinary Terms

Entries will include:

- Mise en place
- Mirepoix
- Bouquet garni
- Roux
- Sauté
- Julienne
- Brunoise
- Confit
- Crème fraîche
- Bain-marie
- Déglacer
- Beurre blanc
- Potage
- Terrine
- Pâté
- Galette
- Pâte brisée

## Chapter 46 — Spanish Culinary Terms

Entries will include:

- Sofrito
- Picada
- Tapas
- Raciones
- Socarrat
- Ajo
- Pimentón
- Alioli
- Tostada
- Empanada
- Pinchos
- Escalivada
- Pisto
- Torrijas
- Arroz

Each entry will follow:

**Original Term — English Meaning — Practical Explanation**

**Estimated length:** 2,500 words

# PART VII — FRENCH AND SPANISH CUISINE IN COMPARISON

## Chapter 47 — Two Culinary Traditions, Many Shared Principles

- Olive oil and butter
- Sauces
- Herbs
- Spices
- Seafood
- Meat
- Rice
- Bread
- Vegetables
- Dining culture
- Regional diversity
- Seasonality
- Preservation
- Communal eating
- Desserts
- Breakfast traditions

This comparison will be respectful and analytical, not competitive. It will show how both cuisines use technique, seasonality, and regional identity to create distinctive food cultures.

**Estimated length:** 2,000 words

# PART VIII — REFERENCE AND LEARNING GUIDES

## Chapter 48 — Conversion Guide

- Grams to ounces
- Kilograms to pounds
- Milliliters to fluid ounces
- Liters to quarts
- Celsius to Fahrenheit
- Teaspoons, tablespoons, and cups
- Common oven-temperature conversions

All conversions will be mathematically checked and rounded consistently.

## Chapter 49 — Seasonal Cooking Guide

### Spring

- Peas
- Asparagus
- Artichokes
- Tender greens
- Fresh herbs

### Summer

- Tomatoes
- Peppers
- Eggplant
- Stone fruit
- Fresh seafood

### Autumn

- Mushrooms
- Squash
- Apples
- Pears
- Root vegetables
- Game and braising cuts

### Winter

- Cabbage
- Beans
- Lentils
- Citrus
- Preserved foods
- Slow-cooked dishes

## Chapter 50 — Beginner’s Learning Path

### Level 1 — Beginner

- Knife skills
- Eggs
- Vegetables
- Basic sauces
- Simple seafood
- Simple desserts

### Level 2 — Intermediate

- Braising
- Stocks
- Complex sauces
- Pastry
- Rice
- Seafood
- Meat

### Level 3 — Advanced

- Laminated dough
- Complex sauces
- Multi-component dishes
- Advanced pastry
- Sophisticated menu planning

## Chapter 51 — Master Recipe Index

Alphabetical index of every recipe, including:

- Original name
- English name where useful
- Cuisine
- Chapter reference

## Chapter 52 — Master Ingredient Index

Alphabetical index of major ingredients and the recipes in which they appear.

**Estimated Part VIII length:** 6,000 words, excluding index entries

# 6. Estimated Word Count by Major Section

| Section | Estimated Words |
| --- | ---: |
| Front matter, introduction, and how to use the book | 4,500 |
| Kitchen foundations | 13,500 |
| French cuisine | 36,000 |
| Spanish cuisine | 32,000 |
| Advanced techniques | 3,500 |
| Menu planning | 2,500 |
| Ingredients and substitutions | 4,000 |
| Culinary dictionary | 2,500 |
| French-Spanish comparison | 2,000 |
| Reference guides and learning path | 4,000 |
| Indexes and navigation material | 2,000 |
| **Estimated total** | **106,500** |

The final manuscript may be shortened during editing if recipe explanations become repetitive or expanded where a technique requires additional teaching.

# 7. Standard Recipe Template

Every recipe will follow the same structure.

# Recipe Name

**Original French or Spanish name, where appropriate**

### Introduction

A concise explanation of:

- What the dish is
- Its culinary significance
- Its country and region
- Why the reader should learn or cook it
- Important variation notes

### Recipe Information

- **Prep Time:**
- **Cook Time:**
- **Total Time:**
- **Servings:**
- **Difficulty:**
- **Cuisine:**
- **Region:**

### Equipment

Only genuinely necessary equipment.

### Ingredients

- Precise quantities
- Metric measurement followed by U.S. equivalent where practical
- Ingredients grouped logically
- Separate sauce, garnish, dough, or component lists where needed

### Instructions

Numbered steps that include:

- Preparation details
- Heat levels
- Visual and sensory cues
- Timing guidance
- Doneness indicators
- Resting instructions

### Chef’s Tips

Two to five practical professional tips.

### Common Mistakes

Likely errors and how to avoid them.

### Variations

Only traditional, widely used, or genuinely useful variations.

### Substitutions

Practical alternatives, with an explanation of how they may affect flavor or texture.

### Serving Suggestions

- Appropriate accompaniments
- Plating guidance
- Course placement
- Beverage suggestions where useful

### Storage

- Refrigeration
- Freezing
- Reheating
- Safe storage duration
- Ingredients that should be stored separately

### Cultural Note

A short regional, historical, or culinary note where appropriate.

### Photography Marker

For selected recipes:

**[PHOTO SUGGESTION: Finished dish]**

or

**[PHOTO SUGGESTION: Step-by-step preparation]**

# 8. Editorial Style Guide

## Voice

The writing should be:

- Warm
- Sophisticated
- Educational
- Encouraging
- Authoritative
- Accessible
- Sensory
- Professional

The reader should feel as though an experienced chef is offering calm, precise guidance in a well-organized kitchen.

## Prose Principles

- Explain the reason behind important actions.
- Prefer useful sensory cues over false precision.
- Avoid filler and exaggerated promotional language.
- Use vivid but controlled descriptions.
- Avoid repetitive transitions and stock phrases.
- Introduce culinary terms in their original language, followed by an English explanation.
- Define unfamiliar ingredients before using them in a recipe.
- Keep historical claims cautious and supportable.
- Distinguish documented history from folklore.
- Never present a modern adaptation as the traditional version.
- Avoid treating France or Spain as culturally uniform.

## Recipe Language

Instructions should:

- Begin with clear action verbs
- Use numbered steps
- State when ingredients are added
- Explain how heat should be adjusted
- Include visual cues such as “deep golden brown,” “translucent,” or “coats the back of a spoon”
- Identify resting periods
- Avoid vague wording such as “cook until done”
- Avoid unnecessary precision when sensory judgment is more reliable

## Terminology and Capitalization

- Use consistent accents in French and Spanish terms.
- Provide translations where useful.
- Use sentence case for most recipe subheadings.
- Capitalize proper dish names consistently.
- Use the same spelling throughout the manuscript.
- Maintain a house style for units, temperatures, and abbreviations.

# 9. Cultural Authenticity Guidelines

1. Identify the country and region of traditional dishes where appropriate.
2. Explain when multiple regional versions exist.
3. Avoid calling one version the only authentic version when culinary traditions vary.
4. Do not combine French and Spanish traditions without clearly identifying the adaptation.
5. Distinguish home cooking, restaurant cooking, classical cuisine, and modern interpretations.
6. Avoid reducing French cuisine to butter, cheese, baguettes, and wine.
7. Avoid reducing Spanish cuisine to paella, tapas, and chorizo.
8. Treat regional identities as fluid and overlapping.
9. Use cautious language for uncertain origins.
10. Do not invent chefs, dates, restaurants, historical events, or origin stories.
11. Explain the cultural role of communal eating, seasonal cooking, and preservation.
12. Do not make claims that a dish is “the original” unless that claim is well established.
13. Describe substitutions honestly rather than calling them identical.
14. Preserve the defining technique and balance of a dish when adapting it for international readers.
15. Use regional terminology accurately and consistently.

# 10. Recipe Quality-Control System

Each recipe will be reviewed against the following checklist:

### Ingredient Check

- Is every ingredient used?
- Are quantities plausible?
- Are ingredients listed in the order of use?
- Are optional ingredients clearly labeled?

### Technique Check

- Are the cooking methods appropriate?
- Are heat levels realistic?
- Are temperatures suitable?
- Are visual and sensory cues included?
- Are resting times specified where necessary?

### Practicality Check

- Can the dish be made in a normal home kitchen?
- Is required equipment listed?
- Are preparation and cooking times realistic?
- Are serving quantities clear?
- Are storage and reheating instructions safe?

### Editorial Check

- Is the recipe distinct from other recipes?
- Does it teach a useful technique?
- Is the regional context accurate?
- Are substitutions sensible?
- Are variations genuinely useful?
- Are instructions logically ordered?

### Safety Check

- Are raw meat, seafood, eggs, and dairy handled safely?
- Are allergen concerns noted where relevant?
- Are unsafe storage claims avoided?
- Are recommended doneness cues appropriate?

# 11. Photography and Design Plan

## Overall Visual Direction

The finished cookbook should feel refined, warm, tactile, and culinary rather than overly decorative. Photography should show real texture, usable portions, natural cooking environments, and the relationship between ingredients and technique.

## Recommended Photography Categories

### Hero Dish Photography

Approximately 55–65 finished-dish images covering:

- Representative French classics
- Representative Spanish classics
- Major seafood dishes
- Signature desserts
- Menu-planning spreads
- Regional specialties

Suggested composition:

- Three-quarter angle for most plated dishes
- Overhead angle for tapas, salads, rice dishes, and shared meals
- Close detail for sauces, pastry layers, crusts, and textures
- Natural props connected to the dish rather than decorative clutter

### Step-by-Step Photography

Approximately 12–16 sequences for techniques that benefit from visual instruction:

- Laminating croissant dough
- Making a roux
- Building a French sauce
- Preparing mirepoix
- Forming gougères
- Making tortilla española
- Building paella
- Achieving socarrat
- Preparing choux pastry
- Folding or shaping empanadas
- Making aioli
- Carving or resting meat

### Ingredient Photography

Approximately 8–12 images:

- Essential French ingredients
- Essential Spanish ingredients
- Fresh herbs and aromatics
- Seafood selection
- Rice varieties
- Pastry ingredients
- Regional pantry ingredients
- Stocks, sauces, and flavor bases

### Regional and Cultural Photography

Approximately 8–10 images:

- Market ingredients
- Shared tables
- Bread and pastry displays
- Tapas service
- Seasonal produce
- Coastal seafood
- French and Spanish kitchen tools
- Hands working with dough or vegetables

## Photography Markers

Markers will appear throughout the manuscript in the following format:

**[PHOTO SUGGESTION: Finished dish — three-quarter view showing the sauce consistency and browned surface, served in a simple ceramic dish with the key herbs visible.]**

**[PHOTO SUGGESTION: Step-by-step preparation — show the progression from raw rice and sofrito to liquid absorption and finished socarrat.]**

**[PHOTO SUGGESTION: Essential French ingredients — butter, shallots, thyme, mushrooms, mustard, and cream arranged on a wooden work surface.]**

## Recommended Layout Considerations

- Use generous space around recipes.
- Keep ingredients and instructions visually distinct.
- Place important technique notes near the recipes that use them.
- Use regional notes as short sidebars or callouts.
- Use consistent icons or labels for difficulty, make-ahead recipes, and storage.
- Avoid placing critical instructions over photographs.
- Ensure digital readers can navigate easily through headings and recipe names.

# 12. Editorial Production Sequence

The manuscript will be produced in the following order:

1. Approve the blueprint.
2. Write and review the front matter.
3. Write the kitchen foundations.
4. Complete the French section chapter by chapter.
5. Complete the Spanish section chapter by chapter.
6. Write the advanced techniques.
7. Create menu-planning chapters.
8. Complete ingredient, substitution, and dictionary sections.
9. Write the comparison chapter.
10. Add conversion, seasonal, and learning guides.
11. Build the recipe and ingredient indexes.
12. Perform a structural, recipe, language, cultural, safety, and publishing audit.
13. Complete a final continuity and consistency edit.

Each cuisine chapter will be reviewed before moving to the next chapter.

## PHASE COMPLETE — READY FOR THE NEXT PHASE

The complete professional blueprint is ready. The next stage is **PHASE 2 — FRONT MATTER**, beginning with the title page, copyright page, dedication, table of contents, introduction, and “How to Use This Book.”
"""


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_inline_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_markdown_document(doc, markdown):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()

        if not line or line == "---":
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                rows = []
                for row in table_lines:
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    if all(re.fullmatch(r":?-+:?", c) for c in cells):
                        continue
                    rows.append(cells)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for r_idx, row in enumerate(rows):
                        for c_idx, value in enumerate(row):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = value
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(9)
                            if r_idx == 0:
                                set_cell_shading(cell, "EADBC8")
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
                doc.add_paragraph()
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", heading_match.group(2))
            p = doc.add_heading(text, level=min(level, 4))
            if level == 1 and "MASTERING" in text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 25, "6B3E26"),
        ("Heading 1", 17, "6B3E26"),
        ("Heading 2", 13, "8A5A3B"),
        ("Heading 3", 11, "A1663F"),
        ("Heading 4", 10.5, "6B3E26"),
    ]:
        style = styles[style_name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "MASTERING THE ART OF FRENCH & SPANISH COOKING  •  PHASE 1 BLUEPRINT"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(115, 115, 115)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(115, 115, 115)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("MASTERING THE ART OF\nFRENCH & SPANISH COOKING")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor(107, 62, 38)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(
        "A Regional, Technique-Driven Guide to the Ingredients, "
        "Traditions, and Recipes of France and Spain"
    )
    run.italic = True
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(138, 90, 59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    run = p.add_run("PHASE 1 — BOOK BLUEPRINT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(95, 95, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    run = p.add_run("By [AUTHOR NAME]")
    run.font.name = "Georgia"
    run.font.size = Pt(12)

    doc.add_page_break()


def add_appendix(doc, source_text):
    doc.add_page_break()
    doc.add_heading("APPENDIX — ORIGINAL MASTER PROMPT", level=1)
    p = doc.add_paragraph(
        "The complete uploaded master prompt is preserved below as the source "
        "brief for the cookbook project."
    )
    p.runs[0].italic = True
    for line in source_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            doc.add_heading(line.lstrip("#").strip(), level=level)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            doc.add_paragraph(line)


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_markdown_document(doc, BLUEPRINT)
    source_text = SOURCE_PROMPT.read_text(encoding="utf-8")
    add_appendix(doc, source_text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()