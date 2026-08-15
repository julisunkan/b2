from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from generate_book_blueprint_docx import configure_document, add_markdown_document


OUTPUT = Path(
    "attached_assets/mastering_french_spanish_cooking_phase2_front_matter.docx"
)


def add_centered_text(doc, text, size=12, bold=False, italic=False, color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_title_page(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(105)
    run = paragraph.add_run("MASTERING THE ART OF\nFRENCH & SPANISH COOKING")
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(107, 62, 38)

    add_centered_text(
        doc,
        "A Regional, Technique-Driven Guide to the Ingredients, "
        "Traditions, and Recipes of France and Spain",
        size=14,
        italic=True,
        color="8A5A3B",
    )
    add_centered_text(doc, "By [AUTHOR NAME]", size=12)
    doc.add_page_break()


def add_copyright_page(doc):
    add_centered_text(doc, "COPYRIGHT", size=17, bold=True, color="6B3E26")
    paragraphs = [
        "Mastering the Art of French & Spanish Cooking",
        "A Regional, Technique-Driven Guide to the Ingredients, Traditions, and Recipes of France and Spain",
        "",
        "Copyright © [YEAR] [AUTHOR NAME]",
        "",
        "All rights reserved.",
        "",
        "No part of this publication may be reproduced, distributed, stored in a retrieval system, or transmitted in any form or by any means, electronic or mechanical, including photocopying, recording, scanning, or other methods, without the prior written permission of the copyright holder, except for brief quotations used in reviews or other uses permitted by applicable copyright law.",
        "",
        "This book is an original educational and culinary work. Recipe names, culinary terminology, and descriptions of established food traditions may necessarily resemble language used in the broader culinary world; the explanations, organization, and instructional material in this volume have been prepared for this publication.",
        "",
        "The information in this book is provided for general educational purposes. Cooking times, temperatures, ingredient quantities, and food-safety guidance may need to be adapted to the reader’s equipment, altitude, ingredients, skill level, dietary requirements, and circumstances. Always use appropriate food-safety practices, including safe handling, storage, and cooking of meat, seafood, eggs, dairy, and other perishable foods.",
        "",
        "Readers with allergies, medical conditions, or special dietary needs should consult an appropriately qualified professional before making changes to their diet. The author and publisher assume no responsibility for injury, loss, or damage resulting from the use or misuse of the information or recipes in this book.",
        "",
        "First edition",
        "",
        "Published by [PUBLISHER NAME]",
        "[CITY, COUNTRY]",
        "",
        "ISBN: [ISBN TO BE ASSIGNED]",
    ]
    for text in paragraphs:
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(9)
        else:
            doc.add_paragraph()
    doc.add_page_break()


def add_dedication(doc):
    add_centered_text(doc, "DEDICATION", size=17, bold=True, color="6B3E26")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(85)
    paragraph.paragraph_format.left_indent = Inches(0.8)
    paragraph.paragraph_format.right_indent = Inches(0.8)
    run = paragraph.add_run(
        "For everyone who has learned that a meal is more than what is placed "
        "on the table—\nfor the hands that teach, the kitchens that welcome, "
        "and the people who gather to share."
    )
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(107, 62, 38)
    doc.add_page_break()


TABLE_OF_CONTENTS = r"""
# CONTENTS

## Front Matter

- Title Page
- Copyright and Disclaimer
- Dedication
- Contents

## Introduction

- The Shared Table of France and Spain

## How to Use This Book

- Reading the Chapters
- Understanding Recipe Information
- Measurements and Temperatures
- Mise en Place
- Adapting Recipes Responsibly

# KITCHEN FOUNDATIONS

## Chapter 1 — Essential Kitchen Equipment

- Core Equipment
- Essential Hand Tools
- Optional Equipment
- Choosing, Maintaining, and Storing Equipment

## Chapter 2 — Knife Skills and Kitchen Safety

- Knife Control
- Fundamental Cuts
- Knife Maintenance
- General Kitchen Safety

## Chapter 3 — Essential Cooking Techniques

- Dry-Heat Techniques
- Moist-Heat Techniques
- Sauce and Flavor Techniques

## Chapter 4 — Flavor Fundamentals

- The Main Components of Flavor
- Browning and the Maillard Reaction
- Balance and Adjustment

## Chapter 5 — Stocks, Broths, and Foundational Preparations

- French Foundations
- Spanish Foundations
- Using Foundations in Later Recipes

## Chapter 6 — Ingredient Fundamentals

- Buying and Storing Ingredients
- Core Ingredient Categories
- Substitution Principles

# PART I — THE ART OF FRENCH COOKING

## Chapter 7 — Understanding French Cuisine

- Culinary Philosophy
- Historical Development
- Regional France

## Chapter 8 — Essential French Ingredients

- Fats and Dairy
- Aromatics and Vegetables
- Herbs and Seasonings
- Proteins and Legumes

## Chapter 9 — The Foundations of French Sauce

- Classic Sauce Families
- Sauce Techniques
- Troubleshooting

## Chapter 10 — French Breakfast and Brunch

## Chapter 11 — French Appetizers and Hors d’Oeuvres

## Chapter 12 — French Soups and Salads

## Chapter 13 — French Poultry and Meat

## Chapter 14 — French Seafood

## Chapter 15 — French Vegetables and Side Dishes

## Chapter 16 — French Breads and Pastries

## Chapter 17 — French Desserts

# PART II — THE ART OF SPANISH COOKING

## Chapter 18 — Understanding Spanish Cuisine

- Culinary History and Influences
- Regional Spain

## Chapter 19 — Essential Spanish Ingredients

- Foundational Ingredients
- Legumes and Preserved Foods
- Seafood and Herbs

## Chapter 20 — Tapas: The Heart of Spanish Dining

- Understanding Tapas
- Building a Balanced Spread

## Chapter 21 — Spanish Soups and Salads

## Chapter 22 — Paella and Spanish Rice Dishes

- Rice Fundamentals
- Regional Rice Traditions

## Chapter 23 — Spanish Seafood

## Chapter 24 — Spanish Meat and Poultry

## Chapter 25 — Spanish Vegetables and Side Dishes

## Chapter 26 — Spanish Breads and Savory Pastries

## Chapter 27 — Spanish Desserts

# PART III — MASTERING THE TECHNIQUES

## Chapter 28 — Heat Control

## Chapter 29 — Building Flavor

## Chapter 30 — Deglazing and Reduction

## Chapter 31 — Emulsions

## Chapter 32 — Working with Seafood

## Chapter 33 — Working with Meat

## Chapter 34 — Cooking Rice

## Chapter 35 — Baking Fundamentals

# PART IV — FRENCH AND SPANISH MENU PLANNING

## Chapter 36 — A Classic French Dinner

## Chapter 37 — A Spanish Tapas Evening

## Chapter 38 — French Dinner for Two

## Chapter 39 — Spanish Family Feast

## Chapter 40 — French Brunch

## Chapter 41 — Mediterranean Seafood Dinner

## Chapter 42 — French and Spanish Celebration Menu

# PART V — INGREDIENTS AND SUBSTITUTIONS

## Chapter 43 — Ingredient Reference

## Chapter 44 — Responsible Substitution

# PART VI — FRENCH AND SPANISH CULINARY DICTIONARY

## Chapter 45 — French Culinary Terms

## Chapter 46 — Spanish Culinary Terms

# PART VII — FRENCH AND SPANISH CUISINE IN COMPARISON

## Chapter 47 — Two Culinary Traditions, Many Shared Principles

# PART VIII — REFERENCE AND LEARNING GUIDES

## Chapter 48 — Conversion Guide

## Chapter 49 — Seasonal Cooking Guide

## Chapter 50 — Beginner’s Learning Path

## Chapter 51 — Master Recipe Index

## Chapter 52 — Master Ingredient Index
"""


INTRODUCTION = r"""
# INTRODUCTION

## The Shared Table of France and Spain

There are many ways to learn a cuisine. You can memorize a collection of recipes, collect unfamiliar ingredients, or follow a set of instructions closely enough to produce a respectable meal. All of these approaches have their place. But the deeper pleasure begins when you understand the ideas beneath the recipes: how heat changes an ingredient, how a sauce gains body, how acidity restores balance, how a region’s climate shapes its pantry, and how a simple dish can carry the memory of a landscape and the habits of generations.

French and Spanish cooking offer two especially rich ways into that understanding. They are distinct culinary traditions, each with its own language, techniques, rhythms, ingredients, and regional identities. Yet they also share a long conversation across the Pyrenees and around the western Mediterranean. Both cuisines value seasonality, careful preparation, good bread, fresh produce, preserved foods, seafood, legumes, and the transformation of modest ingredients into deeply satisfying meals. Both understand that cooking is not only a matter of combining ingredients. It is a matter of attention.

French cooking is often associated with structure. Its classical traditions place great emphasis on foundations: stocks, sauces, pastry, measured technique, and the disciplined control of heat. These associations are useful, but incomplete. French cooking is also the country’s regional home cooking: a pot of beans, a vegetable gratin, a tart made from what is in season, a stew that improves slowly beside the stove, or a simple salad dressed at the last moment. The formal and the familiar belong to the same culinary landscape.

Spanish cooking is often associated with generosity and sociability. A table may fill with small plates, shared rice dishes, grilled seafood, stews, bread rubbed with tomato, olives, peppers, and food designed to be passed from one person to another. This does not mean that Spanish cooking is casual in the sense of being careless. Its apparent simplicity frequently depends on patience: a sofrito cooked until sweet and concentrated, a broth prepared with care, a tortilla turned at the right moment, or rice left undisturbed so that its texture develops properly.

Neither cuisine is a single, uniform system. France contains coastal traditions, mountain cooking, river-valley produce, dairy-rich regions, wine-growing areas, Mediterranean preparations, and culinary traditions shaped by neighboring peoples. Spain is equally regional, with Atlantic and Mediterranean coasts, dry interior plains, northern mountains, fertile river valleys, island cuisines, and communities whose foodways reflect long histories of exchange. The name of a country may appear on a map, but the character of a dish is often found in a smaller place: a town, a coastline, a valley, a market, or a household.

This regional identity is one of the reasons authenticity requires care. A dish can have several traditional forms without one of them canceling the others. Paella, for example, belongs to a broad family of Spanish rice traditions, but its ingredients, pan, method, and cultural meaning are closely connected to particular regions. A French sauce may have a classical name while being adapted in countless home kitchens. A recipe may be traditional in one place, common in another, and modern in a third. This book will identify those distinctions whenever they matter, rather than treating “authentic” as a simple label.

## Why Technique Matters

The recipes in this book are meant to teach more than sequence. They are meant to show cause and effect.

When meat is seared, the goal is not merely to make its surface brown. The heat drives off moisture and creates a layer of concentrated flavor. When onions are sweated slowly, they become softer and sweeter without taking on the deeper flavor of hard browning. When a pan is deglazed, the browned particles left behind become the beginning of a sauce. When a stock is reduced, its flavor and saltiness become more concentrated. When a vinaigrette is emulsified, two liquids that normally separate are brought together in a temporary, stable mixture.

These details change the way a cook works. Instead of relying only on the clock, you begin to notice the sound of a pan, the smell of butter at the edge of browning, the way a sauce clings to a spoon, the resistance of a vegetable under a knife, or the moment when rice has absorbed liquid but still retains its shape. A recipe gives you a route. Technique teaches you how to read the road.

## Ingredients as a Culinary Language

Ingredients do more than provide flavor. They also provide structure, history, and context.

Butter and cream can create richness and carry delicate aromas. Olive oil can add fruitiness, bitterness, and a supple texture. Vinegar and citrus sharpen a dish. Garlic may form a quiet background note or become the dominant voice. Beans and lentils can be humble, substantial, and deeply regional. Seafood reflects geography and season. Bread, rice, potatoes, and pastry reveal how a culture has learned to make nourishment pleasurable.

The best ingredient is not always the rarest one. It is the ingredient whose quality and character suit the dish. A ripe tomato may need little more than salt, oil, and bread. A tough cut of meat may become extraordinary through slow cooking. A sturdy fish may benefit from a simple sauce that gives it contrast rather than competition. Learning to choose ingredients is therefore part of learning to cook.

This book includes substitution guidance because many readers will cook far from France or Spain. A substitute can make a recipe practical, but it should not be described as identical when it is not. The purpose of a substitution is to preserve as much of the dish’s balance and function as possible while acknowledging what changes.

## A Book of Two Traditions

The French and Spanish sections are arranged as parallel culinary studies, but they are not intended to force the cuisines into the same mold. French sauces and Spanish sofritos do not perform exactly the same work. A French pastry and a Spanish empanada may both involve dough, but their textures, purposes, and cultural settings differ. A formal French dinner and a Spanish tapas evening ask different things of the cook and the guests.

At the same time, comparison can be illuminating. Both cuisines ask cooks to understand heat, timing, texture, preservation, and balance. Both use regional ingredients to create food with a strong sense of place. Both can be elaborate or wonderfully direct. A long-simmered stew and a quickly dressed salad may belong to very different parts of a meal, yet each depends on knowing when to stop.

The final comparative section of this book will explore those connections without turning them into a contest. The goal is not to decide which cuisine is better. The goal is to help the reader see more clearly.

## How to Progress Through the Book

You do not need to begin with the most difficult recipe. Start with the dishes that allow you to observe a technique clearly. Learn to handle a knife safely. Cook eggs gently. Sweat onions without rushing them. Make a simple vinaigrette. Roast vegetables until their edges become deeply colored. Prepare a stock when you have time to pay attention to it. These are not minor exercises. They are the vocabulary of a capable kitchen.

From there, move toward sauces, braises, rice, pastry, seafood, and multi-component menus. Repeat recipes when you can. The second attempt often teaches more than the first because you know what to watch for. Keep notes about your pan, oven, ingredients, and timing. A recipe is written for a range of kitchens, but your own kitchen has its particular habits.

Read each recipe completely before beginning. Gather the equipment. Prepare the ingredients. Then cook with your attention on the process rather than on the page. If something looks or smells different from the description, pause and assess it. The instructions are there to support your judgment, not replace it.

## An Invitation to Cook

French and Spanish cooking become less intimidating when they are approached as living traditions rather than tests of worthiness. You do not need a professional kitchen, a perfect collection of pans, or access to every regional ingredient to begin. You need curiosity, a willingness to observe, and enough patience to let a process develop.

Some recipes in this book will ask for precision. Others will teach you when precision gives way to judgment. Together they will show how foundations become sauces, how ingredients become meals, and how meals become occasions for connection.

Begin where you are. Use the ingredients you can find. Respect the dishes you are learning, and allow your own experience to grow through practice. The aim is not simply to reproduce the food of France and Spain. It is to understand enough of their principles that you can cook with confidence, generosity, and care.
"""


HOW_TO_USE = r"""
# HOW TO USE THIS BOOK

## Reading the Chapters

This book is organized as a progression from foundations to application. The opening sections establish the practical knowledge needed in any kitchen: equipment, knife skills, safety, cooking methods, flavor, stocks, and ingredients. The French and Spanish sections then place that knowledge in context, showing how each cuisine builds meals from its own ingredients, techniques, and regional traditions.

You may read the book from beginning to end, especially if you are developing your skills. You may also use it as a reference. A reader preparing a French stew can consult the sections on searing, deglazing, stocks, and braising. Someone attempting paella can study rice, heat management, saffron, stock, and socarrat before beginning. The cross-references are intentional: confident cooking grows when separate techniques begin to connect.

The advanced techniques section is not reserved for professionals. It gathers ideas that appear throughout the recipes and explains them in a concentrated form. Return to it whenever a method feels uncertain. Repetition is part of learning.

## Understanding Recipe Information

Every recipe begins with a short introduction. This identifies the dish, places it within its culinary tradition, and explains why it is useful or enjoyable to learn. Regional notes are included when a dish has a strong connection to a particular place or when several traditional versions exist.

Each recipe also includes the following information:

- **Prep Time** refers to active preparation before or during cooking, including chopping, measuring, trimming, and assembling.
- **Cook Time** refers to time spent on the heat or in the oven. It may not include resting or cooling unless stated.
- **Total Time** includes preparation, cooking, resting, and other essential periods.
- **Servings** provide a practical estimate for the finished dish. Appetizers and tapas may be measured by pieces or portions rather than by full meals.
- **Difficulty** indicates the level of coordination, technique, or timing required. A recipe marked advanced is not necessarily complicated in every respect; it may simply contain a technique that benefits from prior practice.
- **Cuisine** identifies the primary culinary tradition.
- **Region** is included when the dish has a meaningful regional association. It does not imply that every household in that region prepares the dish in exactly the same way.

Recipe times are useful planning tools, not guarantees. The size of your ingredients, the strength of your burner, the material of your pan, the accuracy of your oven, and your own preparation speed will affect the result.

## Measurements and Temperatures

Quantities are given in metric measurements and, where practical, common U.S. equivalents. Metric measurements are generally the more precise reference, particularly for baking, pastry, rice, and ingredients that vary substantially by volume.

When a liquid is listed by volume, the measurement is intended to be practical for a home cook. When flour, sugar, butter, rice, or another ingredient is central to a recipe’s structure, weighing it is often preferable. A cup of flour can vary depending on how it is filled; a measured weight is more consistent.

Oven temperatures are provided in Celsius and Fahrenheit. Ovens vary, and many domestic ovens cycle above and below their stated temperature. For recipes where temperature is especially important, use an oven thermometer if available.

Stovetop heat is described as low, medium-low, medium, medium-high, or high. These terms are starting points rather than universal settings. A heavy pan retains heat differently from a thin pan, and an induction burner behaves differently from a gas flame. Learn to adjust the heat when food is browning too quickly, steaming instead of searing, or reducing more aggressively than intended.

## Mise en Place

The French expression *mise en place* means “put in place.” In practical cooking, it means preparing the ingredients and equipment before the main cooking begins.

Mise en place may include:

- Reading the complete recipe
- Washing and drying produce
- Measuring ingredients
- Chopping vegetables
- Separating eggs
- Preparing stock or broth
- Setting out the correct pan
- Preheating the oven
- Clearing enough workspace

Mise en place is not about making a kitchen look formal. It protects the cook from avoidable interruptions. A sauce can over-reduce while you search for vinegar. Onions can burn while you peel garlic. Pastry can warm while you look for a rolling pin. Preparing first gives you the freedom to watch the food while it cooks.

Mise en place should not be interpreted rigidly. Some ingredients are best cut just before use, and some herbs are best added at the last moment. The point is to understand what must be ready before heat enters the process and what can be prepared during a natural pause.

## How to Read the Instructions

Read all the instructions once before you start. Identify:

1. Which ingredients need advance preparation.
2. Which pans or tools are required.
3. Whether the oven must be preheated.
4. Which steps can happen simultaneously.
5. Where the recipe depends on a visual or sensory cue.
6. Whether the dish needs to rest before serving.

Pay attention to verbs. “Sweat” does not mean brown. “Simmer” does not mean boil. “Reduce” means that liquid volume decreases and flavor becomes more concentrated. “Fold” is different from stirring because it aims to preserve air or structure. Culinary language is concise, but each word carries a practical instruction.

When a recipe describes a sensory cue, treat it as important information. “Until translucent,” “until deeply golden,” “until tender but still firm,” and “until the sauce coats the back of a spoon” are not decorative phrases. They tell you what the food should look, feel, or behave like at the moment you move to the next step.

## Adapting Recipes Responsibly

Recipes are frameworks built around balance. You may need to adapt them because of availability, dietary requirements, equipment, serving size, or personal preference. The best adaptations preserve the role of an ingredient even when the ingredient itself changes.

Before substituting, ask:

- Does the replacement provide a similar flavor?
- Does it contain a similar amount of moisture?
- Will it cook in the same amount of time?
- Does it provide fat, acidity, sweetness, salt, structure, or aroma?
- Will it change the regional identity of the dish?

For example, replacing one leafy herb with another may change aroma but leave the structure of a sauce intact. Replacing a short-grain rice with a long-grain rice may change absorption and texture enough to require a different method. Replacing butter with olive oil may be practical in some vegetable preparations but will change the flavor and behavior of a pastry or sauce.

Substitutions in this book are described honestly. Some are close alternatives; others are compromises that make a dish possible in a different location. Neither is a problem as long as the change is understood.

## Scaling Servings

Many recipes can be doubled, but doubling is not always as simple as multiplying every quantity. A larger amount of food may require a wider pan, a longer reduction time, additional batches, or a different cooking vessel. Crowding a pan can prevent browning. A deep pot can slow evaporation. A larger roast may require a thermometer rather than a simple time calculation.

When reducing a recipe, preserve the proportions of the main ingredients but monitor seasoning carefully. Salt, vinegar, spices, and strong aromatics may not need to be reduced in exactly the same way as the bulk ingredients. Taste and adjust near the end when food safety allows.

## Timing and Preparation

The menu-planning chapters identify what can be made ahead and what should be cooked immediately. In general:

- Stocks, braises, many soups, and some desserts often improve or hold well when prepared in advance.
- Fried foods, delicate seafood, crisp pastry, and certain sauces are usually best close to serving.
- Salads may be prepared in components, with dressing added at the last moment.
- Rice dishes require careful planning because their final texture changes as they rest.
- Pastry benefits from planned chilling and resting periods.

A calm meal is usually the result of preparation rather than speed. Make a written schedule for an ambitious menu, and build in more time than you think you need the first time you cook it.

## Food Safety

Keep raw meat, seafood, eggs, and ready-to-eat foods separate. Wash hands and work surfaces appropriately. Use clean utensils when tasting. Refrigerate perishable food promptly, and do not rely on appearance alone to determine whether food is safe.

Use a reliable thermometer when a recipe calls for a specific internal temperature. For seafood and meat, choose ingredients from reputable sources and follow local food-safety guidance. Be especially careful with shellfish, raw or lightly cooked eggs, poultry, minced meat, and leftovers.

The storage instructions in this book are practical guidance, but local health recommendations and the conditions of your refrigerator should take priority. When in doubt, discard food rather than taking a risk.

## Choosing Where to Begin

If you are new to cooking, begin with a recipe that lets you practice one or two skills without requiring too many components. Eggs, vegetables, simple soups, vinaigrettes, roasted dishes, and straightforward desserts provide useful early lessons.

If you already cook regularly, choose a recipe that exposes a technique you want to improve. You might practice a pan sauce, a braise, a French pastry, a Spanish rice dish, or a shared menu. Do not judge success only by appearance. Ask whether you understood the process more clearly than before.

The most important measure of progress is not whether every meal is perfect. It is whether you can increasingly recognize what is happening in the pan and respond with purpose.
"""


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_copyright_page(doc)
    add_dedication(doc)
    add_markdown_document(doc, TABLE_OF_CONTENTS)
    doc.add_page_break()
    add_markdown_document(doc, INTRODUCTION)
    doc.add_page_break()
    add_markdown_document(doc, HOW_TO_USE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()