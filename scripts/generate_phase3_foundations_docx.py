from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from generate_book_blueprint_docx import configure_document, add_markdown_document


OUTPUT = Path(
    "attached_assets/mastering_french_spanish_cooking_phase3_foundations.docx"
)


def add_centered_text(doc, text, size=12, bold=False, italic=False, color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_title_page(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(105)
    run = paragraph.add_run("MASTERING THE ART OF\nFRENCH & SPANISH COOKING")
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(107, 62, 38)

    add_centered_text(
        doc,
        "A Regional, Technique-Driven Guide to the Ingredients, "
        "Traditions, and Recipes of France and Spain",
        size=14,
        italic=True,
        color="8A5A3B",
    )
    add_centered_text(doc, "PHASE 3 — KITCHEN FOUNDATIONS", size=12, bold=True)
    add_centered_text(doc, "By [AUTHOR NAME]", size=12)
    doc.add_page_break()


FOUNDATIONS = r"""
# PHASE 3 — KITCHEN FOUNDATIONS

## Before the Recipes

French and Spanish cooking become easier to understand when the cook can recognize the work taking place beneath a finished dish. A sauce has a history: perhaps a stock, a browned pan, a reduction, an emulsion, or a carefully controlled roux. A stew begins long before the liquid is added, with knife work, seasoning, browning, and the decision to use gentle heat. A rice dish depends on the shape of the grain, the strength of the broth, the width of the pan, and the moment when the cook stops stirring.

These foundations are not a preliminary obstacle to overcome before the enjoyable cooking begins. They are the enjoyable cooking. They are the skills that allow a cook to move from following directions to making informed decisions.

The chapters in this section cover six areas:

- Choosing and using kitchen equipment
- Handling a knife safely and efficiently
- Understanding the major cooking techniques
- Building and balancing flavor
- Preparing stocks, broths, sauces, and flavor bases
- Selecting, storing, and substituting ingredients

The aim is not to make every cook work in a professional manner. It is to make the home kitchen more legible. When you understand what a pan, a knife, a stock, or an ingredient is doing, you can work with greater confidence and adapt when conditions change.

# CHAPTER 1 — ESSENTIAL KITCHEN EQUIPMENT

## The Right Tool for the Task

Good equipment does not need to be expensive, extensive, or fashionable. The most useful kitchen is one in which the cook can reach the right tool easily, clean it properly, and trust it to behave predictably.

A sharp knife is more useful than a large collection of specialty blades. A well-balanced saucepan is more valuable than several pans that heat unevenly. A sturdy cutting board protects both the cook and the knife. A thermometer can prevent guesswork when cooking meat, poultry, fried foods, sugar, or bread.

The recommendations in this chapter are organized by function. Essential equipment should support the cooking you do regularly. Optional equipment should be added when it solves a real problem rather than because a recipe happens to mention it.

## Core Equipment

### Chef’s Knife

A chef’s knife is the primary tool for chopping, slicing, dicing, mincing, and portioning. A blade between approximately 20 and 25 centimetres (8 and 10 inches) suits many home cooks. The best knife is one that feels secure in the hand and allows the blade to move naturally through food.

Keep the cutting edge sharp. A dull knife requires more pressure, slips more easily, and bruises ingredients. Wash the knife by hand, dry it promptly, and store it where the edge will not strike other utensils.

### Paring Knife

A paring knife is useful for peeling, trimming, removing stems, shaping small vegetables, and making precise cuts. It should complement the chef’s knife rather than replace it.

### Bread Knife

A serrated bread knife cuts through crusty bread without crushing the interior. It is also useful for soft-skinned fruits, ripe tomatoes, and cakes that would compress under a straight blade.

### Cutting Boards

Use a stable board large enough to allow food to be moved away from the cutting area. Wood and durable plastic are both practical materials. Keep separate boards, or wash and sanitize thoroughly, when moving between raw animal products and ready-to-eat foods.

Place a damp towel or nonslip mat beneath a board that moves. A board that slides is a safety problem, not a minor inconvenience.

### Saucepan

A saucepan is designed for liquids, grains, sauces, custards, reheating, and small-batch cooking. A heavy base helps distribute heat and reduces scorching.

### Frying Pan

A frying pan is useful for eggs, sautéed vegetables, fish, cutlets, and pan sauces. The cooking surface should be large enough to leave space between ingredients. Overcrowding traps steam and prevents browning.

### Sauté Pan

A sauté pan has straight, relatively tall sides and a broad base. It is useful when ingredients must be turned in a sauce, when a preparation needs more capacity than a frying pan, or when a dish begins with browning and ends with braising.

### Stockpot

A stockpot should be large enough to hold ingredients and water without filling to the rim. Tall sides reduce evaporation, while a broad enough base allows aromatics or bones to be heated evenly.

### Dutch Oven

A Dutch oven retains heat well and is particularly useful for braising, stewing, baking bread, frying, and cooking dishes that move from stovetop to oven. Its weight is an advantage during long cooking, but it also means that the pot remains hot after the burner is lowered.

### Roasting Pan and Baking Trays

A roasting pan should be strong enough to support the food and allow air to circulate around it. Baking trays are useful for roasting vegetables, baking pastry, drying ingredients, and preparing small dishes.

## Essential Hand Tools

### Whisk

Use a whisk to combine liquids, incorporate air, smooth a roux-based sauce, and create or maintain an emulsion. Choose a shape that reaches into the corners of the pan you use most.

### Wooden Spoon

A wooden spoon is gentle on cookware and useful for stirring sauces, scraping the bottom of a pot, and testing the texture of vegetables. A straight-edged spoon can help clear browned particles from a pan.

### Spatula

A flexible spatula is useful for turning delicate food. A heatproof silicone spatula is useful for scraping bowls and folding mixtures. Keep the tool appropriate to the surface of the pan.

### Tongs

Tongs give the cook control over meat, vegetables, pasta, bread, and seafood without piercing the food unnecessarily. Use tongs with enough length to keep hands away from hot oil or steam.

### Ladle, Strainer, and Colander

A ladle transfers stock and sauce while leaving solids behind. A fine-mesh strainer removes small particles and can help produce a smooth sauce or clear broth. A colander drains pasta, vegetables, and larger ingredients.

### Grater and Peeler

A grater handles cheese, citrus zest, garlic, and firm vegetables. A peeler removes skins efficiently and can also create thin ribbons for salads or garnishes.

### Kitchen Scale and Measuring Tools

A scale provides the most reliable measurements for flour, sugar, rice, butter, and other ingredients whose volume can vary. Measuring spoons remain useful for small quantities of salt, spices, yeast, and extracts.

### Thermometer

An instant-read thermometer provides information that appearance cannot always provide. It is useful for meat, poultry, fish, frying oil, sugar, bread, and reheated foods. Insert it into the thickest part of the food without touching bone or the pan.

## Optional Equipment

### Blender or Food Processor

A blender produces smooth soups, purées, sauces, and emulsions. A food processor is useful for chopping, mixing doughs, making crumbs, and preparing larger quantities. Neither is essential for every kitchen, but both reduce labor for certain preparations.

### Mandoline

A mandoline creates uniform slices and julienne cuts quickly. It must be used with its guard and handled with full attention. A knife is slower but often safer for a small amount of food.

### Mortar and Pestle

A mortar and pestle crush garlic, herbs, nuts, spices, and other ingredients into pastes. This is particularly useful for preparations such as picada and for sauces in which texture matters.

### Stand Mixer and Pastry Tools

A stand mixer can make repeated bread and pastry work easier. A rolling pin, bench scraper, pastry brush, dough cutter, and baking parchment are more fundamental to pastry than the mixer itself.

### Paella Pan and Grill Pan

A paella pan provides a broad cooking surface that supports even rice absorption and surface development. A grill pan can provide seared lines indoors, but it is not a substitute for an outdoor grill and can be more difficult to clean.

## Choosing, Maintaining, and Storing Equipment

Choose equipment based on:

- The type of cooking you do
- The number of people you usually feed
- Storage space
- Ease of cleaning
- Heat source compatibility
- Whether one tool can perform several jobs

Do not place a hot pan directly on a delicate surface. Do not move a hot pan by its lid or by a handle that extends over a burner. Allow cookware to cool before washing when a sudden temperature change could warp it.

Keep knives protected, dry wooden tools fully, and remove food residue before it hardens. A clean, uncluttered workspace makes cooking safer and reduces the temptation to rush.

# CHAPTER 2 — KNIFE SKILLS AND KITCHEN SAFETY

## The Knife Grip

Hold the handle securely with the thumb and forefinger close to the blade, rather than gripping only the end of the handle. The remaining fingers wrap around the handle. This grip gives the cook better control over the blade’s direction and pressure.

The other hand is the guiding hand. Curl the fingertips inward so the knuckles face the blade. Rest the side of the knife against the knuckles and move the guiding hand back gradually as the cut progresses.

Do not cut food that is unstable. Halve a round onion, potato, or squash to create a flat surface before slicing. If an ingredient rocks on the board, stop and make it stable.

## Chopping, Slicing, and Dicing

Chopping creates pieces of a generally similar size without requiring perfect uniformity. It is appropriate for many stocks, stews, and rustic preparations.

Slicing creates flat pieces. The thickness should reflect the cooking method: thin slices cook quickly and may soften or burn easily, while thick slices retain their shape longer.

Dicing creates cubes. Small dice cook quickly and distribute evenly through sauces and soups. Larger dice retain more texture and are useful in roasting and stewing.

## Julienne, Brunoise, and Chiffonade

Julienne means cutting food into thin matchsticks. Brunoise is a very small dice made by first creating uniform strips and then cutting across them. Chiffonade is a technique for leafy herbs or greens: stack the leaves, roll them gently, and slice across the roll into ribbons.

Uniformity matters because pieces of similar size cook at a similar rate. It does not need to be mechanical in a rustic dish, but it becomes important in preparations where texture and appearance depend on even cooking.

## Mincing and Crushing

Mincing produces very small pieces and releases aroma. Garlic can be minced with a knife, crushed with salt, or worked into a paste with a mortar and pestle. The finer the cut, the more quickly the flavor spreads through a dish and the more carefully it must be protected from burning.

## Carving

Allow roasted meat to rest before carving. Identify the direction of the muscle fibers and cut across them when appropriate. Use long, controlled strokes rather than sawing repeatedly. For poultry, separate joints at their natural connections and keep the knife stable.

## Knife Maintenance

A honing steel or ceramic rod can realign an edge that has moved slightly during use; it does not replace sharpening. Sharpening removes a small amount of metal to create a new edge. If a knife struggles to cut a ripe tomato or slips from an onion, it likely needs attention.

Never put a good knife loose in a drawer. Use a blade guard, knife block, magnetic strip installed safely, or another method that protects the edge and the hand.

## Kitchen Safety

Keep handles turned away from the edge of the stove. Lift lids away from the face so steam escapes in the opposite direction. Add wet ingredients to hot oil slowly and from a safe distance. Never leave hot oil unattended.

Wash hands after handling raw meat, poultry, seafood, or eggs. Keep raw ingredients separate from cooked food, use clean utensils for tasting, and chill leftovers promptly. If a surface has contacted raw animal products, clean it before preparing ready-to-eat food.

# CHAPTER 3 — ESSENTIAL COOKING TECHNIQUES

Cooking technique is the controlled transfer of heat, moisture, fat, and time. The same ingredient can become crisp, tender, caramelized, creamy, or dry depending on the method used.

## Sautéing

### What It Is

Sautéing cooks relatively small pieces of food quickly in a shallow amount of fat over medium-high to high heat.

### Why It Works

The hot pan transfers heat rapidly while the fat helps conduct heat across the surface. When the pan is not crowded, surface moisture can evaporate and browning can develop.

### How to Perform It

Preheat the pan, add the fat, and wait until it is hot but not smoking. Add dry ingredients in a single layer. Leave them undisturbed briefly so a crust can form, then turn or stir as needed.

### Common Mistakes

- Adding food to a cold pan
- Crowding the pan
- Using wet ingredients
- Stirring constantly
- Burning garlic or delicate herbs at the beginning

### Signs of Success

Food should have a browned surface where appropriate, remain moist inside, and release easily from the pan.

## Sweating

### What It Is

Sweating cooks aromatics gently in fat without intentionally browning them.

### Why It Works

Gentle heat softens cell walls and releases aroma while preserving a pale color and mild sweetness.

### How to Perform It

Use low to medium-low heat. Add a small amount of salt if appropriate to encourage moisture release. Stir occasionally and lower the heat if the edges begin to color.

### Common Mistakes

- Using too much heat
- Expecting browning
- Cutting pieces unevenly
- Allowing the pan to dry out

### Signs of Success

The aromatics should be soft and translucent, with a sweet smell and little or no browning.

## Searing

### What It Is

Searing exposes the surface of food to high heat to create browning before the food is finished by another method or served.

### Why It Works

Dry surfaces and high heat encourage the development of flavorful browned compounds. Searing also creates contrast between a crust and a tender interior.

### How to Perform It

Dry the food, season appropriately, preheat the pan, and add enough fat to prevent sticking. Place the food down carefully and do not move it until it has released naturally.

### Common Mistakes

- Starting with a wet surface
- Using a pan that is not hot enough
- Overcrowding
- Turning too early
- Burning the fond before the food is finished

### Signs of Success

The surface is deeply golden to brown, not black, and the food has developed a savory roasted aroma.

## Roasting

### What It Is

Roasting cooks food with surrounding dry heat, usually in an oven.

### Why It Works

Dry heat evaporates surface moisture and encourages browning while the interior cooks through.

### How to Perform It

Preheat the oven. Dry and season the ingredients. Leave space between pieces and use a pan that allows air to circulate. Turn ingredients when the recipe calls for even coloring.

### Common Mistakes

- Placing cold, wet food on a crowded tray
- Using a dish too deep for the desired browning
- Opening the oven repeatedly
- Confusing surface color with internal doneness

### Signs of Success

The exterior is appropriately browned and the interior is tender or cooked to the required temperature.

## Baking

### What It Is

Baking uses controlled oven heat to set doughs, batters, custards, pastry, and covered preparations.

### Why It Works

Heat transforms water, starch, protein, fat, sugar, and leavening agents. The structure of the food changes as gases expand, proteins set, and moisture moves.

### How to Perform It

Measure carefully, preheat the oven, use the correct pan, and avoid opening the door unnecessarily during delicate stages. Cool baked goods as directed so their structure can finish setting.

### Common Mistakes

- Measuring flour by packing it tightly
- Using a cold oven
- Substituting pan sizes without adjusting the time
- Overmixing batter
- Cutting bread or pastry before it has set

### Signs of Success

Look for the cues appropriate to the food: a set center, a clean skewer, a crisp crust, a hollow sound, or a temperature reading.

## Grilling and Broiling

Grilling cooks food below or above a strong radiant heat source. Broiling places food close to the heat in an oven. Both methods can produce rapid browning and smoky or roasted flavors.

Dry food, oil lightly when appropriate, and clean the grate or tray. Watch closely because the difference between deep browning and burning can develop quickly. Move thick food to a cooler area or lower heat to finish the interior.

## Braising

### What It Is

Braising cooks a usually firm ingredient first with dry heat and then slowly in a modest amount of liquid.

### Why It Works

Gentle, moist heat softens connective tissue and allows flavor to move into the cooking liquid. Browning contributes depth before the liquid is added.

### How to Perform It

Brown the main ingredient in batches. Build the aromatic base, deglaze, add a measured amount of liquid, cover, and cook slowly until tender.

### Common Mistakes

- Skipping the browning step when it is part of the dish
- Adding too much liquid
- Boiling aggressively
- Failing to skim excess fat
- Serving before the connective tissue has softened

### Signs of Success

The ingredient should yield easily to a fork while remaining intact, and the cooking liquid should taste concentrated and balanced.

## Stewing

Stewing cooks smaller pieces of food slowly in enough liquid to surround them. It is useful for meats, beans, vegetables, and combinations of several ingredients. Keep the heat gentle; a hard boil can break apart delicate ingredients and toughen meat.

## Poaching

Poaching cooks food gently in liquid held below a full simmer. It is suitable for fish, eggs, poultry, fruit, and delicate preparations. The liquid should move slightly but not churn.

## Simmering and Boiling

A simmer has small, occasional bubbles. A boil has vigorous, continuous movement. Use a simmer when you want gradual extraction, gentle cooking, or a clear broth. Use a boil when rapid heat transfer is useful, as with pasta, some vegetables, or sterilizing equipment according to appropriate guidance.

## Blanching

Blanching briefly cooks food in boiling water or steam and often uses an ice bath afterward to stop the cooking. It can set color, soften vegetables, remove skins, or prepare ingredients for another method.

## Frying

Pan-frying uses enough fat to contact part of the food. Deep-frying surrounds the food with hot oil. In both cases, dry surfaces, stable heat, and small batches improve results.

Do not add water to hot oil. Use a thermometer when possible, keep a lid or appropriate fire-safety method nearby, and never leave frying unattended.

## Deglazing

Deglazing adds liquid to a hot pan to dissolve the browned particles left after searing or sautéing. Wine, stock, water, vinegar, or another suitable liquid may be used. Scrape the pan once the liquid loosens the fond, then reduce or use the liquid as the beginning of a sauce.

## Reducing

Reduction removes water through evaporation and concentrates flavor. It also changes saltiness, acidity, and body. Taste before and after reduction; a liquid that seems mild at the beginning may become too salty later.

## Emulsifying

An emulsion combines liquids that naturally separate, usually by slowly dispersing one into another while whisking or blending. Mayonnaise, aioli, hollandaise, vinaigrettes, and butter sauces depend on this principle.

Temperature, rate of addition, agitation, and the presence of an emulsifier all matter. If an emulsion breaks, a small amount of fresh base may help bring it back together.

## Caramelizing

Caramelization is the browning of sugars through heat. It differs from the browning of proteins and sugars together in the Maillard reaction, although both can contribute deep flavor. Sugar can move from pale to bitter quickly; use visual cues and moderate the heat.

## Confit

Confit describes a slow cooking and preservation method in which an ingredient cooks gently in fat or, in some traditions, sugar. The exact technique varies by ingredient and region. Keep temperatures controlled and follow safe storage practices; fat alone does not make a food shelf-stable in a home kitchen.

## Flambéing

Flambéing briefly ignites alcohol in a pan. It is optional rather than essential to most dishes. Use a wide pan, remove it from the heat before adding alcohol when appropriate, keep the face and hands away, and never pour alcohol directly from the bottle into a hot pan or flame.

# CHAPTER 4 — FLAVOR FUNDAMENTALS

## Salt

Salt does more than make food taste salty. In small, appropriate amounts it sharpens aromas, reduces the perception of bitterness, and helps individual flavors become clearer. Add salt in stages when a dish will reduce. Taste near the end before making the final adjustment.

Different salts vary in crystal size and density. A volume measure of fine salt may contain more salt than the same volume of coarse salt. When changing salt types, taste rather than assuming a direct one-to-one replacement.

## Acidity

Acidity gives food brightness and prevents rich preparations from tasting heavy. Vinegar, citrus, wine, tomatoes, cultured dairy, and some fruits contribute acid. Add it gradually. Too much acid can make a dish sharp, but a small amount added at the end can transform a dull sauce or stew.

## Fat

Fat carries aroma, changes texture, and conducts heat. Butter brings dairy richness and can brown into a nutty flavor. Olive oil brings fruitiness and bitterness. Cream softens and rounds. Choose fat according to the role it needs to play, not simply according to habit.

## Sweetness and Bitterness

Sweetness may come from ingredients, browning, reduction, or a small added amount of sugar. Bitterness may come from greens, coffee, char, certain spices, or overcooked aromatics. Both can be useful. Balance does not mean removing every strong flavor; it means giving each flavor a place.

## Umami

Umami is a savory quality found in ingredients such as stocks, aged cheese, mushrooms, tomatoes, cured foods, seafood, and fermented products. It is most effective when layered rather than forced. A good stock, browned vegetables, or a small amount of aged cheese may provide depth without making a dish taste heavy.

## Herbs and Spices

Fresh herbs often provide brightness and aroma. Woody herbs can tolerate longer cooking, while delicate herbs are usually added near the end. Dried herbs are more concentrated and often benefit from early contact with fat or liquid.

Spices should be stored away from heat, light, and moisture. Toasting whole spices can deepen their aroma, but ground spices can burn quickly. Add them at the stage that best suits their purpose.

## Browning and the Maillard Reaction

Browning is one of the most reliable ways to build flavor. Dry the surface of food, heat the pan sufficiently, use a suitable amount of fat, and avoid overcrowding. The browned particles left in the pan are valuable; deglazing turns them into part of the next layer of the dish.

The Maillard reaction involves complex changes between amino compounds and sugars at the surface of food. In practical terms, it is why a browned crust, toasted bread, roasted meat, and deeply colored vegetables taste different from their pale versions. It is not the same as caramelization, which concerns the transformation of sugars.

## Tasting and Adjusting

Taste at useful points:

- After the aromatics soften
- Before and after adding liquid
- Before and after a reduction
- When the main ingredient is nearly cooked
- Immediately before serving

Ask specific questions. Does the dish need salt, acid, richness, sweetness, heat, aroma, or dilution? Do not add several corrections at once. Make one small adjustment, taste again, and allow it to settle.

# CHAPTER 5 — STOCKS, BROTHS, AND FOUNDATIONAL PREPARATIONS

## Stock and Broth

The words “stock” and “broth” are often used interchangeably in home cooking. In this book, stock generally means a liquid built from bones, vegetables, or other ingredients for use as a foundation, while broth may be a more finished liquid intended to be served. The distinction is useful but not absolute.

Good foundational liquids depend on clean ingredients, suitable proportions, controlled heat, and careful straining. A hard boil can make a stock cloudy and emulsify unwanted fat into the liquid. A gentle simmer extracts flavor while keeping the liquid more refined.

Cool cooked liquids quickly in shallow containers. Refrigerate promptly, and use clean utensils when portioning. Refrigerated stock should be used within a short, safe period; freeze portions for longer storage. A layer of solidified fat on a chilled stock can protect it briefly, but it does not make the stock shelf-stable.

## Recipe: Basic Vegetable Stock

### Introduction

This flexible vegetable stock provides a clean, aromatic foundation for soups, sauces, grains, braises, and vegetable dishes. It is intentionally restrained so that it can support both French and Spanish preparations without dominating them.

### Recipe Information

- **Prep Time:** 20 minutes
- **Cook Time:** 45–60 minutes
- **Total Time:** 1 hour 15 minutes
- **Yield:** Approximately 1.5 litres (6 cups)
- **Difficulty:** Beginner
- **Cuisine:** Foundational European preparation

### Equipment

- Stockpot
- Chef’s knife
- Fine-mesh strainer
- Large bowl

### Ingredients

- 2 medium onions, about 300 g (10½ oz), roughly chopped
- 3 medium carrots, about 250 g (9 oz), roughly chopped
- 3 celery stalks, about 180 g (6 oz), roughly chopped
- 1 leek, about 150 g (5 oz), washed and sliced
- 4 garlic cloves, lightly crushed
- 2 bay leaves
- 6 parsley stems
- 6 black peppercorns
- 2.25 litres (9½ cups) cold water

### Instructions

1. Place the onions, carrots, celery, leek, garlic, bay leaves, parsley stems, and peppercorns in a stockpot.
2. Add the cold water. The ingredients should be covered by several centimetres, but the pot should not be filled to the rim.
3. Bring the water slowly to a gentle simmer over medium heat. Do not allow the stock to boil vigorously.
4. Reduce the heat and simmer uncovered for 45–60 minutes, skimming foam only if needed. The vegetables should be soft and the liquid should smell sweet and aromatic.
5. Strain through a fine-mesh strainer. Press the vegetables lightly, but do not mash them if you want a clear stock.
6. Cool quickly in a shallow container. Refrigerate or freeze in useful portions.

### Chef’s Tips

- Avoid strongly colored vegetables if you want a pale stock.
- Do not add salt if the stock will be reduced later.
- A small amount of mushroom can add savoriness, but it will darken the liquid.

### Common Mistakes

- Boiling hard and creating a cloudy stock
- Cooking the vegetables until they become bitter
- Adding too much water and producing a weak liquid

### Storage

Refrigerate promptly and use within 3–4 days. Freeze in sealed containers for approximately 2–3 months for best quality. Leave room for expansion when freezing.

### Cultural Note

Vegetable stocks appear in many culinary traditions, but their exact ingredients vary according to region, season, and purpose. The formula here is a practical foundation rather than a single definitive traditional recipe.

## Recipe: Basic Chicken Stock

### Introduction

Chicken stock provides body and savory depth for soups, sauces, braises, risotto-style preparations, and pan sauces. Bones and joints contribute gelatin, while aromatics provide fragrance.

### Recipe Information

- **Prep Time:** 20 minutes
- **Cook Time:** 3–4 hours
- **Total Time:** About 4 hours 30 minutes
- **Yield:** Approximately 2 litres (8 cups)
- **Difficulty:** Intermediate
- **Cuisine:** French foundation

### Equipment

- Stockpot
- Fine-mesh strainer
- Large bowl
- Ladle

### Ingredients

- 1.5 kg (3¼ lb) chicken wings, backs, necks, or a combination
- 2 medium onions, about 300 g (10½ oz), quartered
- 2 carrots, about 180 g (6 oz), roughly chopped
- 2 celery stalks, about 120 g (4¼ oz), roughly chopped
- 1 leek, about 120 g (4¼ oz), washed and sliced
- 1 garlic head, halved crosswise
- 2 bay leaves
- 8 parsley stems
- 10 black peppercorns
- 3 litres (12¾ cups) cold water

### Instructions

1. Place the chicken pieces in a stockpot and cover with the cold water.
2. Bring the water slowly toward a simmer. As foam rises, skim it from the surface.
3. Add the onions, carrots, celery, leek, garlic, bay leaves, parsley stems, and peppercorns.
4. Maintain a gentle simmer for 3–4 hours. The surface should move slightly rather than boil vigorously.
5. Strain the stock carefully. Do not press the solids if you want a cleaner liquid.
6. Cool quickly. Once chilled, remove excess fat from the surface if desired.

### Chef’s Tips

- Roast the bones and onions first when you want a darker, more robust stock.
- Keep the water level just high enough to cover the ingredients.
- Add fresh water only if necessary; repeated dilution weakens flavor.

### Common Mistakes

- Allowing the stock to boil hard
- Failing to skim excess foam early
- Adding salt before knowing how the stock will be used

### Storage

Refrigerate promptly and use within 3–4 days. Freeze in measured portions for longer storage.

### Cultural Note

Chicken stock is central to many French preparations, but home cooks across Europe make related broths with local ingredients and different levels of concentration.

## Recipe: Basic Beef Stock

### Introduction

Beef stock is a deep, gelatinous foundation for braises, gravies, sauces, and hearty soups. Roasting the bones and vegetables before simmering develops a darker flavor.

### Recipe Information

- **Prep Time:** 25 minutes
- **Cook Time:** 6–8 hours
- **Total Time:** About 8½ hours
- **Yield:** Approximately 2 litres (8 cups)
- **Difficulty:** Intermediate
- **Cuisine:** French foundation

### Equipment

- Roasting pan
- Stockpot
- Fine-mesh strainer
- Large bowl

### Ingredients

- 2.5 kg (5½ lb) beef or veal bones, preferably with joints
- 2 onions, about 300 g (10½ oz), halved
- 3 carrots, about 250 g (9 oz), roughly chopped
- 3 celery stalks, about 180 g (6 oz), roughly chopped
- 2 tablespoons (30 ml) tomato paste
- 3.5 litres (14¾ cups) cold water, or enough to cover
- 2 bay leaves
- 10 parsley stems
- 12 black peppercorns

### Instructions

1. Heat the oven to 220°C (425°F). Arrange the bones in a roasting pan and roast for 35–45 minutes, turning once, until deeply browned.
2. Add the onions, carrots, and celery. Roast for another 20–25 minutes.
3. Transfer the bones and vegetables to a stockpot. Spread the tomato paste over the roasting pan and return it to the oven for 5 minutes, watching carefully so it does not burn.
4. Add a little water to the roasting pan and scrape up the browned particles. Transfer this liquid to the stockpot.
5. Add the remaining cold water, bay leaves, parsley stems, and peppercorns.
6. Bring slowly to a gentle simmer. Cook for 6–8 hours, skimming as needed.
7. Strain, cool quickly, and refrigerate. Remove the solidified fat before using if desired.

### Chef’s Tips

- Deep brown is desirable; blackened bones or tomato paste will make the stock bitter.
- Keep the simmer gentle so the stock does not become greasy and cloudy.
- Reduce the finished stock only after removing excess fat.

### Storage

Refrigerate promptly and use within 3–4 days, or freeze in portions for approximately 2–3 months.

## Recipe: Basic Fish Stock

### Introduction

Fish stock, or fumet-style fish foundation, cooks more quickly than meat stock. It should be aromatic and clean rather than aggressively fishy.

### Recipe Information

- **Prep Time:** 20 minutes
- **Cook Time:** 30–40 minutes
- **Total Time:** 1 hour
- **Yield:** Approximately 1.25 litres (5 cups)
- **Difficulty:** Intermediate
- **Cuisine:** French foundation

### Ingredients

- 1 kg (2¼ lb) fish bones and heads from mild white fish, gills removed
- 1 small onion, about 150 g (5 oz), sliced
- 1 small leek, about 100 g (3½ oz), sliced
- 1 small carrot, about 75 g (2½ oz), sliced
- 100 ml (scant ½ cup) dry white wine
- 1.5 litres (6⅓ cups) cold water
- 1 bay leaf
- 6 parsley stems
- 6 black peppercorns

### Instructions

1. Rinse the bones in cold water. If using heads, remove the gills, which can add bitterness.
2. Place the bones, onion, leek, carrot, bay leaf, parsley stems, and peppercorns in a pot.
3. Add the wine and water. Bring slowly to a gentle simmer.
4. Cook for 25–30 minutes, skimming as necessary. Do not cook fish bones for many hours; they can make the liquid bitter.
5. Strain carefully and cool promptly.

### Common Mistakes

- Using oily fish when a clean, mild stock is needed
- Leaving gills attached to fish heads
- Boiling hard
- Overcooking the bones

### Storage

Refrigerate promptly and use within 1–2 days, or freeze in small portions.

## Recipe: Court-Bouillon

### Introduction

Court-bouillon is an aromatic poaching liquid used for fish, shellfish, poultry, and vegetables. It is usually lightly acidic and is meant to flavor the food gently rather than become a finished sauce.

### Recipe Information

- **Prep Time:** 10 minutes
- **Cook Time:** 25 minutes
- **Total Time:** 35 minutes
- **Yield:** Approximately 1.5 litres (6 cups)
- **Difficulty:** Beginner
- **Cuisine:** French foundation

### Ingredients

- 1.5 litres (6⅓ cups) water
- 250 ml (1 cup) dry white wine or 60 ml (¼ cup) white wine vinegar plus additional water
- 1 onion, sliced
- 1 carrot, sliced
- 1 celery stalk, sliced
- 1 bay leaf
- 6 parsley stems
- 6 black peppercorns
- 1 teaspoon fine salt

### Instructions

1. Combine all ingredients in a pot.
2. Bring to a simmer and cook for 20–25 minutes.
3. Taste and adjust the seasoning. Use warm for poaching or cool according to the recipe.

### Cultural Note

The term describes a family of French aromatic liquids rather than one fixed formula. Wine, vinegar, herbs, vegetables, and spices may vary by the food being poached.

## Recipe: Mirepoix

### Introduction

Mirepoix is a French aromatic mixture traditionally based on onion, carrot, and celery. It is a flavor base rather than a finished dish.

### Recipe Information

- **Prep Time:** 10 minutes
- **Cook Time:** Varies by recipe
- **Yield:** Approximately 600 g (1¼ lb)
- **Difficulty:** Beginner
- **Cuisine:** French foundation

### Ingredients

- 300 g (10½ oz) onion, diced
- 150 g (5¼ oz) carrot, diced
- 150 g (5¼ oz) celery, diced

### Instructions

1. Cut all vegetables into pieces of similar size.
2. Sweat them gently in butter or oil when a pale, sweet foundation is desired.
3. Brown them more deeply when building a darker stock, braise, or sauce.

### Chef’s Tips

- Use a fine dice for quick-cooking sauces and a larger cut for long-simmered stocks.
- Leek can be used alongside or in place of some onion.
- Do not salt heavily if the preparation will be reduced.

## Recipe: Bouquet Garni

### Introduction

A bouquet garni is a bundle of herbs used to perfume stocks, sauces, soups, and stews. It is removed before serving.

### Ingredients

- 2 parsley stems
- 1 small bay leaf
- 2 sprigs thyme
- Kitchen twine

### Instructions

1. Wrap the herbs together in a small piece of leek green or tie them directly with kitchen twine.
2. Add the bundle to the pot at the beginning of cooking.
3. Remove before blending or serving.

### Cultural Note

Bouquet garni formulas vary. The principle is more important than a rigid list: use a restrained bundle of herbs that complements the dish and can be removed easily.

## Recipe: Basic Spanish Sofrito

### Introduction

Sofrito is a slowly cooked aromatic foundation found in many Spanish preparations. Its exact composition varies by region and dish. This version emphasizes onion, garlic, tomato, and olive oil, cooked until concentrated and sweet.

### Recipe Information

- **Prep Time:** 15 minutes
- **Cook Time:** 45–60 minutes
- **Total Time:** About 1 hour 15 minutes
- **Yield:** Approximately 500 g (about 2 cups)
- **Difficulty:** Beginner
- **Cuisine:** Spanish foundation

### Ingredients

- 60 ml (¼ cup) extra-virgin olive oil
- 2 medium onions, about 300 g (10½ oz), finely chopped
- 1 red bell pepper, about 150 g (5 oz), finely chopped
- 4 garlic cloves, finely chopped
- 600 g (21 oz) ripe tomatoes, grated or finely crushed
- 1 teaspoon fine salt
- ½ teaspoon sweet or smoked paprika, optional

### Instructions

1. Warm the olive oil in a wide pan over medium-low heat.
2. Add the onions and pepper with the salt. Cook slowly for 15–20 minutes, stirring occasionally, until soft and sweet.
3. Add the garlic and paprika, if using. Cook for 1 minute.
4. Add the tomatoes. Cook uncovered for 25–35 minutes, stirring occasionally, until the mixture is thick, glossy, and no longer watery.
5. Taste and adjust the salt. Use immediately or cool promptly.

### Chef’s Tips

- A wide pan encourages evaporation and concentration.
- Do not rush the onion stage; raw sharpness should disappear before the tomato is added.
- Sofrito can be left slightly textured or blended smooth according to its later use.

### Common Mistakes

- Using high heat and burning the garlic
- Leaving excess water in the finished mixture
- Adding too much paprika and obscuring the tomato

### Storage

Refrigerate promptly and use within 3–4 days. Freeze in small portions for approximately 2–3 months.

### Cultural Note

Spanish sofrito is not a single standardized sauce. Ingredients and proportions shift across regions and households. Its defining idea is the slow development of aromatics into a concentrated base.

## Recipe: Spanish Seafood Broth

### Introduction

This quick broth is designed for rice dishes, seafood stews, sauces, and soups. Shells, fish bones, aromatics, and paprika create a savory liquid with a distinctly Spanish direction.

### Recipe Information

- **Prep Time:** 20 minutes
- **Cook Time:** 35 minutes
- **Total Time:** 55 minutes
- **Yield:** Approximately 1.5 litres (6 cups)
- **Difficulty:** Intermediate
- **Cuisine:** Spanish-inspired foundation

### Ingredients

- 500 g (1 lb 2 oz) shrimp or prawn shells and heads, if available
- 500 g (1 lb 2 oz) mild fish bones
- 2 tablespoons (30 ml) olive oil
- 1 onion, chopped
- 1 small carrot, chopped
- 1 celery stalk, chopped
- 3 garlic cloves, crushed
- 1 teaspoon smoked paprika
- 100 ml (scant ½ cup) dry white wine
- 1.75 litres (7½ cups) cold water
- 1 bay leaf
- 4 parsley stems

### Instructions

1. Heat the olive oil in a stockpot over medium heat. Add the shells and cook for 3–4 minutes, stirring, until fragrant and lightly colored.
2. Add the onion, carrot, celery, and garlic. Cook for 5 minutes.
3. Stir in the paprika for 30 seconds, taking care not to burn it.
4. Add the wine and scrape the bottom of the pot.
5. Add the water, bay leaf, parsley stems, and fish bones. Bring to a gentle simmer.
6. Cook for 25 minutes, skimming as needed. Strain carefully and cool promptly.

### Common Mistakes

- Burning paprika
- Boiling the broth hard
- Cooking fish bones for too long
- Leaving shell fragments in the strained liquid

### Storage

Refrigerate promptly and use within 1–2 days, or freeze in measured portions.

## Recipe: Picada

### Introduction

Picada is a Catalan-style mixture used to finish or thicken certain stews, sauces, and rice dishes. Its ingredients vary, but nuts, bread, garlic, herbs, and sometimes saffron are common elements.

### Recipe Information

- **Prep Time:** 10 minutes
- **Cook Time:** 5 minutes
- **Total Time:** 15 minutes
- **Yield:** Approximately 120 g (½ cup)
- **Difficulty:** Beginner
- **Cuisine:** Catalan foundation

### Ingredients

- 40 g (1½ oz) blanched almonds
- 1 small slice country bread, about 30 g (1 oz)
- 2 garlic cloves
- 2 tablespoons chopped parsley
- 2 tablespoons (30 ml) olive oil
- 2 tablespoons water or cooking liquid
- Small pinch saffron, optional

### Instructions

1. Warm the olive oil in a small pan. Fry the bread and almonds gently until golden. Add the garlic during the final minute.
2. Transfer to a mortar or food processor. Add the parsley and saffron, if using.
3. Pound or blend to a coarse paste, adding the water or cooking liquid gradually.
4. Stir into a stew or sauce near the end of cooking, allowing several minutes for the flavors to integrate.

### Cultural Note

Picada is a technique and family of preparations rather than one fixed recipe. Its purpose is to add body, aroma, and texture to the dish in which it is used.

# CHAPTER 6 — INGREDIENT FUNDAMENTALS

## Buying Ingredients

Buy for the dish you intend to cook, not for an imagined future kitchen. Freshness matters, but freshness has different signs for different ingredients. A crisp herb, a firm fish, a fragrant spice, and a mature cheese each communicate quality in different ways.

When shopping, consider:

- Whether the ingredient is in season
- Whether it will be used quickly
- Whether the package protects it from air, light, and moisture
- Whether a less expensive ingredient can perform the same culinary function
- Whether the recipe depends on a particular variety

## Seasonality

Seasonal ingredients often have better flavor and may be more affordable. They also connect a dish to the time of year in which it traditionally makes sense. Seasonality does not mean refusing all preserved or imported ingredients; drying, salting, fermenting, curing, freezing, and bottling are central to both French and Spanish cooking.

Use the seasonal guide later in the book to choose appropriate substitutions when a named ingredient is unavailable.

## Storage

Keep a dry pantry cool, dark, and organized. Store spices away from the stove. Refrigerate perishable ingredients promptly, and do not wash delicate produce until you are ready to use it unless the item’s storage method calls for washing.

Store herbs according to their type. Tender herbs may last longer with their stems in water and a loose covering. Woody herbs generally prefer a dry, cool environment. Keep onions and potatoes separate in a cool, dark, ventilated place rather than in a sealed plastic bag.

## Fats

Butter contributes dairy flavor and can be browned. Olive oil contributes fruitiness and is central to many Spanish preparations as well as parts of southern French cooking. Cream adds body. Duck fat, pork fat, and other animal fats contribute distinct flavor and may be traditional in particular regions.

Choose a fat for:

- Flavor
- Heat tolerance
- Texture
- Cultural context
- Whether it will be served or removed

## Aromatics

Onions, shallots, garlic, leeks, celery, carrots, peppers, and herbs create the first layer of many dishes. Their size and cooking stage matter. A finely minced aromatic dissolves into a sauce; a large piece perfumes a stock and can be removed.

Do not brown every aromatic automatically. A pale, sweet base and a deeply browned base serve different purposes.

## Grains and Rice

Rice varies in length, starch, absorption, and texture. Short- and medium-grain varieties can release more starch and create a creamier result. Some Spanish rice dishes depend on grains that absorb flavor while remaining distinct.

Read the recipe before substituting rice. A different grain may require different liquid, heat, timing, and stirring.

## Legumes

Beans, chickpeas, and lentils provide substance and absorb the flavors around them. Dried legumes require planning. Sort and rinse them, soak when appropriate, and cook until tender before adding acidic ingredients if the particular legume tends to remain firm in acid.

Canned legumes are practical. Rinse them when the recipe calls for a clean flavor, and reserve their liquid only when its texture and salt level suit the preparation.

## Dairy and Eggs

Cream, butter, cheese, yogurt, crème fraîche, and eggs provide richness, structure, emulsification, and color. Temperature matters. Cold butter can mount a sauce; softened butter creams into dough; eggs can thicken a custard but may scramble if heated too aggressively.

Introduce hot liquid gradually to eggs when tempering. Keep dairy-based sauces below a hard boil unless the recipe specifically requires otherwise.

## Flour, Bread, and Pastry Ingredients

Flour develops gluten when hydrated and worked. The desired amount of gluten depends on the food: bread needs strength, while tender pastry needs restraint. Measure accurately and avoid adding extra flour simply because dough feels sticky at the beginning; time and rest can change its texture.

Bread serves as food, thickener, coating, vehicle, and ingredient. Stale bread is valuable in crumbs, soups, picada, and desserts.

## Meat and Poultry

Choose a cut according to the cooking method. Tender cuts suit quick cooking. Tougher cuts with connective tissue often become excellent through slow braising. Dry the surface before searing, season appropriately, and allow the finished meat to rest when the cut and method require it.

Keep raw meat cold, separate from ready-to-eat foods, and use a reliable thermometer when doneness cannot be judged safely by appearance.

## Seafood

Fresh seafood should smell clean and briny rather than strongly fishy. Flesh should be firm and moist. Shellfish should be purchased from a reputable source and handled according to local food-safety guidance.

Seafood is often overcooked because the cook relies on time rather than observation. Watch for changes in translucency, firmness, and flaking. Carryover cooking continues after seafood leaves the heat.

## Herbs and Spices

Fresh herbs provide high notes and should often be added near the end. Woody herbs can tolerate longer cooking. Dried herbs need time to hydrate and release flavor. Spices should be tasted periodically for freshness; old spices may be safe but dull.

Use herbs and spices to reinforce a dish’s identity. More is not always more. A restrained amount that arrives at the correct moment is often more effective than a large amount added without a plan.

## Substitution Principles

When an ingredient is unavailable, identify its function before choosing a replacement.

### If the ingredient provides acidity

Consider lemon juice, wine vinegar, sherry vinegar, wine, tomatoes, or cultured dairy, depending on the dish. Each changes the flavor, so add gradually.

### If the ingredient provides fat

Consider butter, olive oil, cream, or another suitable fat. Account for differences in water content, flavor, and heat behavior.

### If the ingredient provides body

Consider reduction, a roux, puréed vegetables, bread, ground nuts, legumes, or a starch, while recognizing that each creates a different texture.

### If the ingredient provides aroma

Choose a related herb, spice, aromatic, or preserved ingredient. Use a smaller amount first and taste.

### If the ingredient provides structure

Be cautious. Replacing rice, flour, eggs, gelatin, or a leavening agent can change the entire recipe. Look for a tested alternative rather than making a casual swap.

## A Practical Pantry

A useful French and Spanish pantry may include:

- Extra-virgin olive oil
- Unsalted butter
- Vinegars
- Dried beans and lentils
- Several rice varieties
- Flour
- Sugar
- Canned tomatoes
- Stock ingredients
- Garlic and onions
- Dried herbs and spices
- Mustard
- Olives
- Nuts and almonds
- Anchovies or other preserved seafood
- Saffron or a suitable regional seasoning

The pantry should support cooking, not become a museum of unused ingredients. Buy small quantities of items used rarely, label opened packages, and replace spices when their aroma fades.

## Closing Note

The foundations in this section are meant to be practiced repeatedly. Chop vegetables for a soup and notice whether your cuts are even. Make stock and observe how a simmer differs from a boil. Taste a sauce before and after reduction. Roast the same vegetable at two different temperatures. These small comparisons build the judgment that recipes cannot provide by themselves.

Once the kitchen becomes more understandable, the recipes become more generous. You can see what they are asking, recognize when a result is developing correctly, and make a useful adjustment when your ingredients or equipment differ from the page.

**PHASE COMPLETE — READY FOR THE NEXT PHASE**
"""


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_markdown_document(doc, FOUNDATIONS)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()